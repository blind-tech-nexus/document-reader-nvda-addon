import gui
import wx
import ui
import webbrowser
import threading
import os
import sys
import json
import time
import re
import subprocess
import tempfile
import asyncio
import urllib.parse
import urllib.request
from scriptHandler import script
import globalPluginHandler
import globalVars
import addonHandler
from logHandler import log
from . import languages

base_path = os.path.dirname(__file__)
libs_path = os.path.join(base_path, "libs")
if libs_path not in sys.path:
    sys.path.insert(0, libs_path)

import secrets

try:
    import fitz
except ImportError:
    raise RuntimeError("PyMuPDF not found in libs directory")

try:
    import edge_tts
except ImportError as e:
    edge_tts = None
    log.warning(f"edge_tts import failed: {e}")

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    import docx
except ImportError:
    docx = None

try:
    import pytesseract
except ImportError as e:
    log.warning(f"pyteseract import failed: {e}")
    pytesseract = None

try:
    from PIL import Image, ImageEnhance, ImageFilter
except ImportError as e:
    log.warning(f"pillow import failed: {e}")
    Image = None
    ImageEnhance = None
    ImageFilter = None

try:
    from playsound import playsound
except ImportError:
    playsound = None

try:
    from deep_translator import GoogleTranslator
except ImportError:
    GoogleTranslator = None

try:
    import ebooklib
    from ebooklib import epub
except ImportError:
    ebooklib = None
    epub = None

try:
    from bs4 import BeautifulSoup
except ImportError:
    BeautifulSoup = None

addonHandler.initTranslation()

DB_PATH = os.path.join(os.path.dirname(__file__), "doc_reader_data.json")

def show_msg(msg, title=_("Information"), is_error=False, silent=False):
    ui.message(msg)
    if not silent:
        style = wx.ICON_ERROR if is_error else wx.ICON_INFORMATION
        wx.MessageBox(msg, title, wx.OK | style)

def thread_safe_msg(msg, title=_("Information"), is_error=False, silent=False):
    wx.CallAfter(show_msg, msg, title, is_error, silent)

def load_data():
    if not os.path.exists(DB_PATH):
        return {
            "bookmarks": [], "notes": [], "history": {}, "recent_files": {},
            "settings": {}, "annotations": [], "highlights": [],
            "reading_sessions": [], "custom_stamps": [], "doc_tags": {},
            "page_labels": {}, "extraction_profiles": [], "tts_cache": {},
            "search_history": [], "split_views": {},
            "saved_voices": []
        }
    try:
        with open(DB_PATH, "r", encoding="utf-8") as f:
            return json.load(f)
    except:
        return {
            "bookmarks": [], "notes": [], "history": {}, "recent_files": {},
            "settings": {}, "annotations": [], "highlights": [],
            "reading_sessions": [], "custom_stamps": [], "doc_tags": {},
            "page_labels": {}, "extraction_profiles": [], "tts_cache": {},
            "search_history": [], "split_views": {},
            "saved_voices": []
        }

def save_data(data):
    try:
        with open(DB_PATH, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
    except Exception as e:
        log.error(f"Failed to save data: {e}")

class ProcessingDialog(wx.Dialog):
    def __init__(self, parent, title=_("Processing...")):
        super().__init__(parent, title=title, style=wx.CAPTION)
        self.SetSize((350, 130))
        self.Centre()
        self._destroyed = False
        main_sizer = wx.BoxSizer(wx.VERTICAL)
        self.static_text = wx.StaticText(self, label=_("Please wait..."))
        main_sizer.Add(self.static_text, 0, wx.ALL | wx.CENTER, 20)
        self.progress = wx.Gauge(self, range=100)
        main_sizer.Add(self.progress, 0, wx.ALL | wx.EXPAND, 10)
        self.SetSizer(main_sizer)
        self.static_text.SetFocus()
        self.Bind(wx.EVT_CLOSE, self.on_close)

    def on_close(self, event):
        self._destroyed = True
        self.Destroy()

    def update(self, text, value=None):
        if self._destroyed:
            return
        wx.CallAfter(self.static_text.SetLabel, text)
        if value is not None:
            wx.CallAfter(self.progress.SetValue, value)

class ResumeDialog(wx.Dialog):
    def __init__(self, parent, last_page, total_pages):
        super().__init__(parent, title=_("Resume Reading"))
        self.SetSize((400, 180))
        self.Centre()
        self.last_page = last_page
        main_sizer = wx.BoxSizer(wx.VERTICAL)
        msg = _("You were on page {last} of {total}. Resume from there?").format(last=last_page+1, total=total_pages)
        label = wx.StaticText(self, label=msg)
        label.Wrap(350)
        main_sizer.Add(label, 0, wx.ALL, 10)
        btn_sizer = wx.BoxSizer(wx.HORIZONTAL)
        resume_btn = wx.Button(self, label=_("&Resume"))
        resume_btn.Bind(wx.EVT_BUTTON, self.on_resume)
        btn_sizer.Add(resume_btn, 0, wx.RIGHT, 10)
        start_btn = wx.Button(self, label=_("Start From &Beginning"))
        start_btn.Bind(wx.EVT_BUTTON, self.on_start)
        btn_sizer.Add(start_btn, 0, wx.RIGHT, 10)
        cancel_btn = wx.Button(self, label=_("&Cancel"))
        cancel_btn.Bind(wx.EVT_BUTTON, self.on_cancel)
        btn_sizer.Add(cancel_btn, 0)
        main_sizer.Add(btn_sizer, 0, wx.ALL | wx.CENTER, 10)
        self.SetSizer(main_sizer)
        resume_btn.SetFocus()

    def on_resume(self, event):
        self.EndModal(wx.ID_OK)

    def on_start(self, event):
        self.EndModal(wx.ID_NO)

    def on_cancel(self, event):
        self.EndModal(wx.ID_CANCEL)

class PasswordPromptDialog(wx.Dialog):
    def __init__(self, parent):
        super().__init__(parent, title=_("PDF Password"))
        self.SetSize((350, 180))
        self.Centre()
        main_sizer = wx.BoxSizer(wx.VERTICAL)
        label = wx.StaticText(self, label=_("Enter the PDF password:"))
        main_sizer.Add(label, 0, wx.ALL | wx.CENTER, 10)
        self.pwd_ctrl_hidden = wx.TextCtrl(self, style=wx.TE_PASSWORD)
        self.pwd_ctrl_visible = wx.TextCtrl(self)
        self.pwd_ctrl_visible.Hide()
        main_sizer.Add(self.pwd_ctrl_hidden, 0, wx.ALL | wx.EXPAND, 10)
        main_sizer.Add(self.pwd_ctrl_visible, 0, wx.ALL | wx.EXPAND, 10)
        self.show_cb = wx.CheckBox(self, label=_("&Show password"))
        self.show_cb.Bind(wx.EVT_CHECKBOX, self.on_show)
        main_sizer.Add(self.show_cb, 0, wx.ALL, 5)
        btn_sizer = wx.BoxSizer(wx.HORIZONTAL)
        ok_btn = wx.Button(self, label=_("&OK"))
        ok_btn.Bind(wx.EVT_BUTTON, self.on_ok)
        btn_sizer.Add(ok_btn, 0, wx.RIGHT, 10)
        cancel_btn = wx.Button(self, label=_("&Cancel"))
        cancel_btn.Bind(wx.EVT_BUTTON, self.on_cancel)
        btn_sizer.Add(cancel_btn, 0)
        main_sizer.Add(btn_sizer, 0, wx.ALL | wx.CENTER, 10)
        self.SetSizer(main_sizer)
        self.pwd_ctrl_hidden.SetFocus()

    def on_show(self, event):
        if self.show_cb.IsChecked():
            self.pwd_ctrl_visible.SetValue(self.pwd_ctrl_hidden.GetValue())
            self.pwd_ctrl_hidden.Hide()
            self.pwd_ctrl_visible.Show()
            self.pwd_ctrl_visible.SetFocus()
        else:
            self.pwd_ctrl_hidden.SetValue(self.pwd_ctrl_visible.GetValue())
            self.pwd_ctrl_visible.Hide()
            self.pwd_ctrl_hidden.Show()
            self.pwd_ctrl_hidden.SetFocus()
        self.Layout()

    def get_password(self):
        if self.show_cb.IsChecked():
            return self.pwd_ctrl_visible.GetValue()
        return self.pwd_ctrl_hidden.GetValue()

    def on_ok(self, event):
        self.EndModal(wx.ID_OK)

    def on_cancel(self, event):
        self.EndModal(wx.ID_CANCEL)

class ProtectPdfDialog(wx.Dialog):
    def __init__(self, parent):
        super().__init__(parent, title=_("Protect PDF"))
        self.SetSize((380, 250))
        self.Centre()
        main_sizer = wx.BoxSizer(wx.VERTICAL)
        pwd_label = wx.StaticText(self, label=_("New password:"))
        main_sizer.Add(pwd_label, 0, wx.ALL, 5)
        self.pwd_ctrl_hidden = wx.TextCtrl(self, style=wx.TE_PASSWORD)
        self.pwd_ctrl_visible = wx.TextCtrl(self)
        self.pwd_ctrl_visible.Hide()
        main_sizer.Add(self.pwd_ctrl_hidden, 0, wx.ALL | wx.EXPAND, 10)
        main_sizer.Add(self.pwd_ctrl_visible, 0, wx.ALL | wx.EXPAND, 10)
        confirm_label = wx.StaticText(self, label=_("Confirm new password:"))
        main_sizer.Add(confirm_label, 0, wx.ALL, 5)
        self.confirm_ctrl_hidden = wx.TextCtrl(self, style=wx.TE_PASSWORD)
        self.confirm_ctrl_visible = wx.TextCtrl(self)
        self.confirm_ctrl_visible.Hide()
        main_sizer.Add(self.confirm_ctrl_hidden, 0, wx.ALL | wx.EXPAND, 10)
        main_sizer.Add(self.confirm_ctrl_visible, 0, wx.ALL | wx.EXPAND, 10)
        self.show_cb = wx.CheckBox(self, label=_("&Show password"))
        self.show_cb.Bind(wx.EVT_CHECKBOX, self.on_show)
        main_sizer.Add(self.show_cb, 0, wx.ALL, 5)
        btn_sizer = wx.BoxSizer(wx.HORIZONTAL)
        protect_btn = wx.Button(self, label=_("&Confirm Protect"))
        protect_btn.Bind(wx.EVT_BUTTON, self.on_protect)
        btn_sizer.Add(protect_btn, 0, wx.RIGHT, 10)
        cancel_btn = wx.Button(self, label=_("&Cancel"))
        cancel_btn.Bind(wx.EVT_BUTTON, self.on_cancel)
        btn_sizer.Add(cancel_btn, 0)
        main_sizer.Add(btn_sizer, 0, wx.ALL | wx.CENTER, 10)
        self.SetSizer(main_sizer)
        self.pwd_ctrl_hidden.SetFocus()

    def on_show(self, event):
        show = self.show_cb.IsChecked()
        if show:
            self.pwd_ctrl_visible.SetValue(self.pwd_ctrl_hidden.GetValue())
            self.confirm_ctrl_visible.SetValue(self.confirm_ctrl_hidden.GetValue())
            self.pwd_ctrl_hidden.Hide()
            self.confirm_ctrl_hidden.Hide()
            self.pwd_ctrl_visible.Show()
            self.confirm_ctrl_visible.Show()
            self.pwd_ctrl_visible.SetFocus()
        else:
            self.pwd_ctrl_hidden.SetValue(self.pwd_ctrl_visible.GetValue())
            self.confirm_ctrl_hidden.SetValue(self.confirm_ctrl_visible.GetValue())
            self.pwd_ctrl_visible.Hide()
            self.confirm_ctrl_visible.Hide()
            self.pwd_ctrl_hidden.Show()
            self.confirm_ctrl_hidden.Show()
            self.pwd_ctrl_hidden.SetFocus()
        self.Layout()

    def on_protect(self, event):
        pwd = self.pwd_ctrl_visible.GetValue() if self.show_cb.IsChecked() else self.pwd_ctrl_hidden.GetValue()
        confirm = self.confirm_ctrl_visible.GetValue() if self.show_cb.IsChecked() else self.confirm_ctrl_hidden.GetValue()
        if pwd != confirm:
            wx.MessageBox(_("Passwords do not match."), _("Error"), wx.OK | wx.ICON_ERROR)
            return
        self.EndModal(wx.ID_OK)

    def on_cancel(self, event):
        self.EndModal(wx.ID_CANCEL)

    def get_password(self):
        if self.show_cb.IsChecked():
            return self.pwd_ctrl_visible.GetValue()
        return self.pwd_ctrl_hidden.GetValue()

class DocumentViewerDialog(wx.Dialog):
    def __init__(self, parent, file_path):
        super().__init__(parent, title=_("Document Reader Panel"), style=wx.DEFAULT_DIALOG_STYLE | wx.RESIZE_BORDER | wx.MAXIMIZE_BOX)
        self.file_path = file_path
        self.file_type = None
        self.pdf_doc = None
        self.docx_doc = None
        self.epub_book = None
        self.sections = []
        self.total_sections = 0
        self.current_section = 0
        self.reading_mode = "text"
        self.search_results = []
        self.current_search_index = -1
        self.ocr_cache = {}
        self.zoom_level = 100
        self.contrast_mode = False
        self.invert_colors = False
        self.dark_mode = False
        self.text_wrap = True
        self.auto_scroll = False
        self.auto_scroll_timer = None
        self.read_aloud = False
        self.current_highlight = None
        self.page_cache = {}
        self.extraction_profile = "default"
        self.split_view_active = False
        self.font_size = 12
        self.line_spacing = 1.0
        self.reading_speed = 0
        self.force_reload = False
        self.settings = {}
        self.saved_voices = []
        self.ffplay_process = None
        self.current_audio_file = None
        self.translated_sections = {}
        self.original_texts = {}
        self.SetSize((900, 750))
        self.Centre()
        self.InitUI()
        self.load_settings()
        self.load_document()
        self.Bind(wx.EVT_CLOSE, self.on_close_dialog)

    def InitUI(self):
        main_sizer = wx.BoxSizer(wx.VERTICAL)

        toolbar_sizer = wx.BoxSizer(wx.HORIZONTAL)
        file_btn = wx.Button(self, label=_("&File"))
        file_btn.Bind(wx.EVT_BUTTON, self.on_file_menu)
        toolbar_sizer.Add(file_btn, 0, wx.ALL, 2)
        edit_btn = wx.Button(self, label=_("&Edit"))
        edit_btn.Bind(wx.EVT_BUTTON, self.on_edit_menu)
        toolbar_sizer.Add(edit_btn, 0, wx.ALL, 2)
        view_btn = wx.Button(self, label=_("&View"))
        view_btn.Bind(wx.EVT_BUTTON, self.on_view_menu)
        toolbar_sizer.Add(view_btn, 0, wx.ALL, 2)
        tools_btn = wx.Button(self, label=_("&Tools"))
        tools_btn.Bind(wx.EVT_BUTTON, self.on_tools_menu)
        toolbar_sizer.Add(tools_btn, 0, wx.ALL, 2)
        nav_btn = wx.Button(self, label=_("&Navigate"))
        nav_btn.Bind(wx.EVT_BUTTON, self.on_nav_menu)
        toolbar_sizer.Add(nav_btn, 0, wx.ALL, 2)
        manipulations_btn = wx.Button(self, label=_("&Manipulations"))
        manipulations_btn.Bind(wx.EVT_BUTTON, self.on_manipulations_menu)
        toolbar_sizer.Add(manipulations_btn, 0, wx.ALL, 2)
        settings_btn = wx.Button(self, label=_("&Settings"))
        settings_btn.Bind(wx.EVT_BUTTON, self.on_settings_dialog)
        toolbar_sizer.Add(settings_btn, 0, wx.ALL, 2)
        help_btn = wx.Button(self, label=_("&Help"))
        help_btn.Bind(wx.EVT_BUTTON, self.on_help_menu)
        toolbar_sizer.Add(help_btn, 0, wx.ALL, 2)

        self.close_panel_btn = wx.Button(self, label=_("Close Reader Panel"))
        self.close_panel_btn.Bind(wx.EVT_BUTTON, self.on_close_dialog)
        toolbar_sizer.Add(self.close_panel_btn, 0, wx.ALL, 2)

        main_sizer.Add(toolbar_sizer, 0, wx.EXPAND)

        info_sizer = wx.BoxSizer(wx.HORIZONTAL)
        self.toolbar_text = wx.StaticText(self, label=_("Section {} of {}").format(0, 0))
        info_sizer.Add(self.toolbar_text, 0, wx.ALL | wx.CENTER, 5)
        info_sizer.AddStretchSpacer()
        self.reading_mode_label = wx.StaticText(self, label=_("Mode:"))
        info_sizer.Add(self.reading_mode_label, 0, wx.RIGHT | wx.CENTER, 5)
        self.mode_choice = wx.Choice(self, choices=[_("Plain Text"), _("Formatted")])
        self.mode_choice.SetSelection(0)
        self.mode_choice.Bind(wx.EVT_CHOICE, self.on_mode_change)
        info_sizer.Add(self.mode_choice, 0, wx.RIGHT, 10)
        main_sizer.Add(info_sizer, 0, wx.ALL | wx.EXPAND, 5)

        content_sizer = wx.BoxSizer(wx.HORIZONTAL)
        self.text_ctrl = wx.TextCtrl(self, style=wx.TE_MULTILINE | wx.TE_READONLY | wx.TE_RICH2 | wx.TE_NOHIDESEL)
        self.text_ctrl.SetFont(wx.Font(self.font_size, wx.FONTFAMILY_TELETYPE, wx.FONTSTYLE_NORMAL, wx.FONTWEIGHT_NORMAL))
        content_sizer.Add(self.text_ctrl, 1, wx.ALL | wx.EXPAND, 5)
        main_sizer.Add(content_sizer, 1, wx.EXPAND, 5)

        nav_sizer = wx.BoxSizer(wx.HORIZONTAL)
        self.prev_btn = wx.Button(self, label=_("◀ &Prev"))
        self.prev_btn.Bind(wx.EVT_BUTTON, self.on_prev)
        nav_sizer.Add(self.prev_btn, 0, wx.RIGHT, 5)
        self.section_choice = wx.Choice(self, choices=[])
        self.section_choice.Bind(wx.EVT_CHOICE, self.on_section_choice)
        nav_sizer.Add(self.section_choice, 1, wx.RIGHT, 5)
        self.next_btn = wx.Button(self, label=_("&Next ▶"))
        self.next_btn.Bind(wx.EVT_BUTTON, self.on_next)
        nav_sizer.Add(self.next_btn, 0, wx.RIGHT, 10)
        self.go_btn = wx.Button(self, label=_("&Go"))
        self.go_btn.Bind(wx.EVT_BUTTON, self.on_go_to_section)
        nav_sizer.Add(self.go_btn, 0)
        main_sizer.Add(nav_sizer, 0, wx.ALL | wx.EXPAND, 5)

        action_sizer = wx.BoxSizer(wx.HORIZONTAL)
        self.search_btn = wx.Button(self, label=_("&Search"))
        self.search_btn.Bind(wx.EVT_BUTTON, self.on_search_dialog)
        action_sizer.Add(self.search_btn, 0, wx.RIGHT, 5)
        self.bm_add_btn = wx.Button(self, label=_("&Bookmark"))
        self.bm_add_btn.Bind(wx.EVT_BUTTON, self.on_add_bookmark)
        action_sizer.Add(self.bm_add_btn, 0, wx.RIGHT, 5)
        self.toc_btn = wx.Button(self, label=_("&TOC"))
        self.toc_btn.Bind(wx.EVT_BUTTON, self.on_toc)
        action_sizer.Add(self.toc_btn, 0, wx.RIGHT, 5)
        self.ocr_btn = wx.Button(self, label=_("&OCR"))
        self.ocr_btn.Bind(wx.EVT_BUTTON, self.on_ocr)
        action_sizer.Add(self.ocr_btn, 0, wx.RIGHT, 5)
        self.export_btn = wx.Button(self, label=_("E&xport"))
        self.export_btn.Bind(wx.EVT_BUTTON, self.on_export_menu)
        action_sizer.Add(self.export_btn, 0, wx.RIGHT, 5)
        self.note_btn = wx.Button(self, label=_("&Note"))
        self.note_btn.Bind(wx.EVT_BUTTON, self.on_add_note)
        action_sizer.Add(self.note_btn, 0, wx.RIGHT, 5)
        self.read_aloud_btn = wx.Button(self, label=_("&Read"))
        self.read_aloud_btn.Bind(wx.EVT_BUTTON, self.on_read_aloud)
        action_sizer.Add(self.read_aloud_btn, 0)
        main_sizer.Add(action_sizer, 0, wx.ALL | wx.CENTER, 5)

        status_sizer = wx.BoxSizer(wx.HORIZONTAL)
        self.status_bar = wx.StaticText(self, label="")
        status_sizer.Add(self.status_bar, 1, wx.ALL | wx.CENTER, 5)
        self.position_label = wx.StaticText(self, label="")
        status_sizer.Add(self.position_label, 0, wx.ALL | wx.CENTER, 5)
        main_sizer.Add(status_sizer, 0, wx.ALL | wx.EXPAND, 5)

        self.SetSizer(main_sizer)
        self._setup_accelerators()
        self.text_ctrl.Bind(wx.EVT_KEY_DOWN, self.on_text_key_down)

    def _create_menu_from_defs(self, defs):
        menu = wx.Menu()
        for label, id_, handler in defs:
            item = menu.Append(id_, label)
            self.Bind(wx.EVT_MENU, handler, id=item.GetId())
        return menu

    def _setup_accelerators(self):
        self.accel_ids = {
            'import': wx.NewIdRef(),
            'export': wx.NewIdRef(),
            'properties': wx.NewIdRef(),
            'close': wx.NewIdRef(),
            'search': wx.NewIdRef(),
            'find_next': wx.NewIdRef(),
            'find_previous': wx.NewIdRef(),
            'copy_page': wx.NewIdRef(),
            'copy_all': wx.NewIdRef(),
            'zoom_in': wx.NewIdRef(),
            'zoom_out': wx.NewIdRef(),
            'zoom_reset': wx.NewIdRef(),
            'high_contrast': wx.NewIdRef(),
            'invert_colors': wx.NewIdRef(),
            'split_view': wx.NewIdRef(),
            'ocr': wx.NewIdRef(),
            'extraction_profile': wx.NewIdRef(),
            'read_aloud': wx.NewIdRef(),
            'auto_scroll': wx.NewIdRef(),
            'compare_sections': wx.NewIdRef(),
            'statistics': wx.NewIdRef(),
            'add_bookmark': wx.NewIdRef(),
            'bookmarks_manager': wx.NewIdRef(),
            'toc': wx.NewIdRef(),
            'annotations': wx.NewIdRef(),
            'highlights': wx.NewIdRef(),
            'go_to_section': wx.NewIdRef(),
            'prev_section': wx.NewIdRef(),
            'next_section': wx.NewIdRef(),
            'help': wx.NewIdRef(),
            'about': wx.NewIdRef(),
            'menu_navigator': wx.NewIdRef(),
            'highlight_selection': wx.NewIdRef(),
            'translate_section': wx.NewIdRef(),
        }
        self.Bind(wx.EVT_MENU, self.on_import, id=self.accel_ids['import'])
        self.Bind(wx.EVT_MENU, self.on_export_menu, id=self.accel_ids['export'])
        self.Bind(wx.EVT_MENU, self.on_properties, id=self.accel_ids['properties'])
        self.Bind(wx.EVT_MENU, self.on_close_dialog, id=self.accel_ids['close'])
        self.Bind(wx.EVT_MENU, self.on_search_dialog, id=self.accel_ids['search'])
        self.Bind(wx.EVT_MENU, self.on_find_next, id=self.accel_ids['find_next'])
        self.Bind(wx.EVT_MENU, self.on_find_previous, id=self.accel_ids['find_previous'])
        self.Bind(wx.EVT_MENU, self.on_copy_page, id=self.accel_ids['copy_page'])
        self.Bind(wx.EVT_MENU, self.on_copy_all, id=self.accel_ids['copy_all'])
        self.Bind(wx.EVT_MENU, self.on_zoom_in, id=self.accel_ids['zoom_in'])
        self.Bind(wx.EVT_MENU, self.on_zoom_out, id=self.accel_ids['zoom_out'])
        self.Bind(wx.EVT_MENU, self.on_zoom_reset, id=self.accel_ids['zoom_reset'])
        self.Bind(wx.EVT_MENU, self.on_high_contrast, id=self.accel_ids['high_contrast'])
        self.Bind(wx.EVT_MENU, self.on_invert_colors, id=self.accel_ids['invert_colors'])
        self.Bind(wx.EVT_MENU, self.on_split_view, id=self.accel_ids['split_view'])
        self.Bind(wx.EVT_MENU, self.on_ocr, id=self.accel_ids['ocr'])
        self.Bind(wx.EVT_MENU, self.on_extraction_profile, id=self.accel_ids['extraction_profile'])
        self.Bind(wx.EVT_MENU, self.on_read_aloud, id=self.accel_ids['read_aloud'])
        self.Bind(wx.EVT_MENU, self.on_auto_scroll, id=self.accel_ids['auto_scroll'])
        self.Bind(wx.EVT_MENU, self.on_compare_sections, id=self.accel_ids['compare_sections'])
        self.Bind(wx.EVT_MENU, self.on_statistics, id=self.accel_ids['statistics'])
        self.Bind(wx.EVT_MENU, self.on_add_bookmark, id=self.accel_ids['add_bookmark'])
        self.Bind(wx.EVT_MENU, self.on_bookmarks_manager, id=self.accel_ids['bookmarks_manager'])
        self.Bind(wx.EVT_MENU, self.on_toc, id=self.accel_ids['toc'])
        self.Bind(wx.EVT_MENU, self.on_annotations, id=self.accel_ids['annotations'])
        self.Bind(wx.EVT_MENU, self.on_highlights, id=self.accel_ids['highlights'])
        self.Bind(wx.EVT_MENU, self.on_go_to_section, id=self.accel_ids['go_to_section'])
        self.Bind(wx.EVT_MENU, self.on_prev, id=self.accel_ids['prev_section'])
        self.Bind(wx.EVT_MENU, self.on_next, id=self.accel_ids['next_section'])
        self.Bind(wx.EVT_MENU, self.on_help, id=self.accel_ids['help'])
        self.Bind(wx.EVT_MENU, self.on_about, id=self.accel_ids['about'])
        self.Bind(wx.EVT_MENU, self.show_menu_navigator, id=self.accel_ids['menu_navigator'])
        self.Bind(wx.EVT_MENU, self.on_highlight_selection, id=self.accel_ids['highlight_selection'])
        self.Bind(wx.EVT_MENU, self.on_translate_section, id=self.accel_ids['translate_section'])
        accel_entries = [
            (wx.ACCEL_CTRL, ord('I'), self.accel_ids['import']),
            (wx.ACCEL_CTRL | wx.ACCEL_SHIFT, ord('E'), self.accel_ids['export']),
            (wx.ACCEL_CTRL, ord('M'), self.accel_ids['properties']),
            (wx.ACCEL_CTRL, ord('W'), self.accel_ids['close']),
            (wx.ACCEL_CTRL, ord('F'), self.accel_ids['search']),
            (wx.ACCEL_NORMAL, wx.WXK_F3, self.accel_ids['find_next']),
            (wx.ACCEL_SHIFT, wx.WXK_F3, self.accel_ids['find_previous']),
            (wx.ACCEL_CTRL, ord('C'), self.accel_ids['copy_page']),
            (wx.ACCEL_CTRL | wx.ACCEL_SHIFT, ord('C'), self.accel_ids['copy_all']),
            (wx.ACCEL_CTRL, ord('+'), self.accel_ids['zoom_in']),
            (wx.ACCEL_CTRL, ord('-'), self.accel_ids['zoom_out']),
            (wx.ACCEL_CTRL, ord('0'), self.accel_ids['zoom_reset']),
            (wx.ACCEL_CTRL, ord('H'), self.accel_ids['high_contrast']),
            (wx.ACCEL_CTRL | wx.ACCEL_SHIFT, ord('I'), self.accel_ids['invert_colors']),
            (wx.ACCEL_CTRL | wx.ACCEL_SHIFT, ord('V'), self.accel_ids['split_view']),
            (wx.ACCEL_CTRL | wx.ACCEL_SHIFT, ord('O'), self.accel_ids['ocr']),
            (wx.ACCEL_CTRL | wx.ACCEL_SHIFT, ord('P'), self.accel_ids['extraction_profile']),
            (wx.ACCEL_CTRL, ord('R'), self.accel_ids['read_aloud']),
            (wx.ACCEL_CTRL | wx.ACCEL_SHIFT, ord('A'), self.accel_ids['auto_scroll']),
            (wx.ACCEL_CTRL | wx.ACCEL_SHIFT, ord('C'), self.accel_ids['compare_sections']),
            (wx.ACCEL_CTRL | wx.ACCEL_SHIFT, ord('S'), self.accel_ids['statistics']),
            (wx.ACCEL_CTRL, ord('B'), self.accel_ids['add_bookmark']),
            (wx.ACCEL_CTRL | wx.ACCEL_SHIFT, ord('B'), self.accel_ids['bookmarks_manager']),
            (wx.ACCEL_CTRL, ord('T'), self.accel_ids['toc']),
            (wx.ACCEL_CTRL | wx.ACCEL_SHIFT, ord('N'), self.accel_ids['annotations']),
            (wx.ACCEL_CTRL | wx.ACCEL_SHIFT, ord('H'), self.accel_ids['highlights']),
            (wx.ACCEL_CTRL, ord('G'), self.accel_ids['go_to_section']),
            (wx.ACCEL_ALT, wx.WXK_PAGEUP, self.accel_ids['prev_section']),
            (wx.ACCEL_ALT, wx.WXK_PAGEDOWN, self.accel_ids['next_section']),
            (wx.ACCEL_NORMAL, wx.WXK_F1, self.accel_ids['help']),
            (wx.ACCEL_CTRL, wx.WXK_F1, self.accel_ids['about']),
            (wx.ACCEL_CTRL | wx.ACCEL_SHIFT, ord('M'), self.accel_ids['menu_navigator']),
            (wx.ACCEL_CTRL | wx.ACCEL_SHIFT, ord('L'), self.accel_ids['highlight_selection']),
            (wx.ACCEL_CTRL | wx.ACCEL_SHIFT, ord('T'), self.accel_ids['translate_section']),
        ]
        self.SetAcceleratorTable(wx.AcceleratorTable(accel_entries))

    def _file_menu_defs(self):
        return [
            (_("&Import Document\tCtrl+I"), wx.ID_ANY, self.on_import),
            (_("&Export\tCtrl+Shift+E"), wx.ID_ANY, self.on_export_menu),
            (_("P&roperties\tCtrl+M"), wx.ID_ANY, self.on_properties),
            (_("Protect PDF"), wx.ID_ANY, self.on_protect_pdf),
            (_("&Close\tCtrl+W"), wx.ID_ANY, self.on_close_dialog),
        ]

    def _edit_menu_defs(self):
        return [
            (_("&Find\tCtrl+F"), wx.ID_ANY, self.on_search_dialog),
            (_("Find &Next\tF3"), wx.ID_ANY, self.on_find_next),
            (_("Find &Previous\tShift+F3"), wx.ID_ANY, self.on_find_previous),
            (_("Find and Replace"), wx.ID_ANY, self.on_find_replace),
            (_("&Copy Section Text\tCtrl+C"), wx.ID_ANY, self.on_copy_page),
            (_("Copy &All Text\tCtrl+Shift+C"), wx.ID_ANY, self.on_copy_all),
            (_("Highlight &Selection\tCtrl+Shift+L"), wx.ID_ANY, self.on_highlight_selection),
        ]

    def _view_menu_defs(self):
        return [
            (_("&Zoom In\tCtrl++"), wx.ID_ANY, self.on_zoom_in),
            (_("Zoom &Out\tCtrl+-"), wx.ID_ANY, self.on_zoom_out),
            (_("&Reset Zoom\tCtrl+0"), wx.ID_ANY, self.on_zoom_reset),
            (_("&High Contrast\tCtrl+H"), wx.ID_ANY, self.on_high_contrast),
            (_("&Invert Colors\tCtrl+Shift+I"), wx.ID_ANY, self.on_invert_colors),
            (_("Dark Mode"), wx.ID_ANY, self.on_dark_mode),
            (_("&Split View\tCtrl+Shift+V"), wx.ID_ANY, self.on_split_view),
        ]

    def _tools_menu_defs(self):
        translate_label = _("Reverse to original language") if self.translated_sections.get(self.current_section, False) else _("Translate Section\tCtrl+Shift+T")
        return [
            (_("&OCR\tCtrl+Shift+O"), wx.ID_ANY, self.on_ocr),
            (_("&Text Extraction Profile\tCtrl+Shift+P"), wx.ID_ANY, self.on_extraction_profile),
            (_("&Read Aloud\tCtrl+R"), wx.ID_ANY, self.on_read_aloud),
            (_("&Auto Scroll\tCtrl+Shift+A"), wx.ID_ANY, self.on_auto_scroll),
            (_("&Compare Sections\tCtrl+Shift+C"), wx.ID_ANY, self.on_compare_sections),
            (_("&Statistics\tCtrl+Shift+S"), wx.ID_ANY, self.on_statistics),
            (translate_label, wx.ID_ANY, self.on_translate_section),
            (_("Extract Links"), wx.ID_ANY, self.on_extract_links),
            (_("Voice Preview"), wx.ID_ANY, self.on_voice_preview),
            (_("Redact Text"), wx.ID_ANY, self.on_redact_text),
        ]

    def _nav_menu_defs(self):
        return [
            (_("&Bookmarks\tCtrl+B"), wx.ID_ANY, self.on_add_bookmark),
            (_("Bookmarks &Manager\tCtrl+Shift+B"), wx.ID_ANY, self.on_bookmarks_manager),
            (_("Export Bookmarks"), wx.ID_ANY, self.on_export_bookmarks),
            (_("&Table of Contents\tCtrl+T"), wx.ID_ANY, self.on_toc),
            (_("&Annotations\tCtrl+Shift+N"), wx.ID_ANY, self.on_annotations),
            (_("&Highlights\tCtrl+Shift+H"), wx.ID_ANY, self.on_highlights),
            (_("&Go To Section\tCtrl+G"), wx.ID_ANY, self.on_go_to_section),
            (_("&Previous Section\tAlt+PageUp"), wx.ID_ANY, self.on_prev),
            (_("&Next Section\tAlt+PageDown"), wx.ID_ANY, self.on_next),
        ]

    def _manipulations_menu_defs(self):
        return [
            (_("&Add Annotation"), wx.ID_ANY, self.on_add_annotation),
            (_("&Rotate Page Clockwise"), wx.ID_ANY, self.on_rotate_cw),
            (_("Rotate Page Counter&clockwise"), wx.ID_ANY, self.on_rotate_ccw),
            (_("&Delete Current Section"), wx.ID_ANY, self.on_delete_page),
            (_("E&xtract Sections"), wx.ID_ANY, self.on_extract_pages),
            (_("&Merge PDF"), wx.ID_ANY, self.on_merge_pdf),
            (_("S&plit PDF"), wx.ID_ANY, self.on_split_pdf),
            (_("Extract Images"), wx.ID_ANY, self.on_extract_images),
            (_("Save Section as Image"), wx.ID_ANY, self.on_save_page_image),
            (_("Add Watermark"), wx.ID_ANY, self.on_add_watermark),
        ]

    def _help_menu_defs(self):
        return [
            (_("&Keyboard Shortcuts\tF1"), wx.ID_ANY, self.on_help),
            (_("&About\tCtrl+F1"), wx.ID_ANY, self.on_about),
        ]

    def on_file_menu(self, event):
        self.PopupMenu(self._create_menu_from_defs(self._file_menu_defs()))
    def on_edit_menu(self, event):
        self.PopupMenu(self._create_menu_from_defs(self._edit_menu_defs()))
    def on_view_menu(self, event):
        self.PopupMenu(self._create_menu_from_defs(self._view_menu_defs()))
    def on_tools_menu(self, event):
        self.PopupMenu(self._create_menu_from_defs(self._tools_menu_defs()))
    def on_nav_menu(self, event):
        self.PopupMenu(self._create_menu_from_defs(self._nav_menu_defs()))
    def on_manipulations_menu(self, event):
        self.PopupMenu(self._create_menu_from_defs(self._manipulations_menu_defs()))
    def on_help_menu(self, event):
        self.PopupMenu(self._create_menu_from_defs(self._help_menu_defs()))

    def on_settings_dialog(self, event):
        dlg = SettingsDialog(self)
        if dlg.ShowModal() == wx.ID_OK:
            self.apply_settings(dlg.get_settings())
        dlg.Destroy()

    def apply_settings(self, settings):
        self.settings.update(settings)
        if "zoom" in settings:
            self.zoom_level = settings["zoom"]
        if "contrast" in settings:
            self.contrast_mode = settings["contrast"]
        if "invert" in settings:
            self.invert_colors = settings["invert"]
        if "font_size" in settings:
            self.font_size = settings["font_size"]
        if "reading_speed" in settings:
            self.reading_speed = settings["reading_speed"]
        if "extraction_profile" in settings:
            self.extraction_profile = settings["extraction_profile"]
            self.page_cache.clear()
        if "voice" in settings:
            self.settings["voice"] = settings["voice"]
        if "tts_speed" in settings:
            self.settings["tts_speed"] = settings["tts_speed"]
        if "tts_pitch" in settings:
            self.settings["tts_pitch"] = settings["tts_pitch"]
        self.update_text_display()
        self.load_section(self.current_section)
        self.save_settings()
        show_msg(_("Settings applied."), _("Success"))

    def load_settings(self):
        data = load_data()
        self.settings = data.get("settings", {})
        self.zoom_level = self.settings.get("zoom", 100)
        self.contrast_mode = self.settings.get("contrast", False)
        self.invert_colors = self.settings.get("invert", False)
        self.font_size = self.settings.get("font_size", 12)
        self.line_spacing = self.settings.get("line_spacing", 1.0)
        self.reading_speed = self.settings.get("reading_speed", 0)
        self.extraction_profile = self.settings.get("extraction_profile", "default")
        self.saved_voices = data.get("saved_voices", [])
        self.update_text_display()

    def save_settings(self):
        data = load_data()
        data["settings"] = {
            "zoom": self.zoom_level,
            "contrast": self.contrast_mode,
            "invert": self.invert_colors,
            "font_size": self.font_size,
            "line_spacing": self.line_spacing,
            "reading_speed": self.reading_speed,
            "reading_mode": self.reading_mode,
            "extraction_profile": self.extraction_profile,
            "voice": self.settings.get("voice", "en-US-AriaNeural"),
            "tts_speed": self.settings.get("tts_speed", "+0%"),
            "tts_pitch": self.settings.get("tts_pitch", "+0Hz"),
            "export_dir": self.settings.get("export_dir", ""),
            "tts_volume": self.settings.get("tts_volume", 100),
            "tts_chunking": self.settings.get("tts_chunking", True),
        }
        data["saved_voices"] = self.saved_voices
        save_data(data)

    def detect_file_type(self):
        ext = os.path.splitext(self.file_path)[1].lower()
        if ext == '.pdf':
            return 'pdf'
        elif ext == '.docx':
            return 'docx'
        elif ext == '.epub':
            return 'epub'
        else:
            return None

    def load_document(self):
        self.file_type = self.detect_file_type()
        if not self.file_type:
            show_msg(_("Unsupported file format."), _("Error"), True)
            self.Close()
            return

        if self.file_type == 'pdf':
            self.load_pdf()
        elif self.file_type == 'docx':
            self.load_docx()
        elif self.file_type == 'epub':
            self.load_epub()
        else:
            show_msg(_("Unknown file type."), _("Error"), True)
            self.Close()
            return

    def load_pdf(self):
        try:
            self.pdf_doc = fitz.open(self.file_path)
            if self.pdf_doc.needs_pass:
                dlg = PasswordPromptDialog(self)
                if dlg.ShowModal() == wx.ID_OK:
                    pw = dlg.get_password()
                    dlg.Destroy()
                    if not self.pdf_doc.authenticate(pw):
                        show_msg(_("Incorrect password."), _("Error"), True)
                        self.Close()
                        return
                else:
                    dlg.Destroy()
                    self.Close()
                    return
        except Exception as e:
            show_msg(_("Unable to open PDF: {}").format(str(e)), _("Error"), True)
            self.Close()
            return
        self.total_sections = len(self.pdf_doc)
        self.sections = [{'type': 'pdf_page', 'page': i} for i in range(self.total_sections)]
        self.populate_mode_choices_pdf()
        self.section_choice.Clear()
        self.section_choice.AppendItems([str(i+1) for i in range(self.total_sections)])
        data = load_data()
        history = data.get("history", {})
        last_section = history.get(self.file_path, 0)
        if last_section > 0 and last_section < self.total_sections:
            resume_dlg = ResumeDialog(self, last_section, self.total_sections)
            res = resume_dlg.ShowModal()
            resume_dlg.Destroy()
            if res == wx.ID_OK:
                self.load_section(last_section)
            elif res == wx.ID_NO:
                self.load_section(0)
            else:
                self.load_section(0)
        else:
            self.load_section(0)
        self.update_toolbar()
        self.save_recent_file()

    def populate_mode_choices_pdf(self):
        modes = [_("Blocks"), _("Text"), _("Words"), _("HTML"), _("Structured")]
        self.mode_choice.SetItems(modes)
        self.mode_choice.SetSelection(1)

    def load_docx(self):
        if docx is None:
            show_msg(_("python-docx library not found."), _("Error"), True)
            self.Close()
            return
        try:
            self.docx_doc = docx.Document(self.file_path)
        except Exception as e:
            show_msg(_("Unable to open DOCX: {}").format(str(e)), _("Error"), True)
            self.Close()
            return
        self.sections = []
        current_heading = None
        current_text = []

        for element in self.docx_doc.element.body:
            if element.tag.endswith('p'):
                para = docx.text.paragraph.Paragraph(element, self.docx_doc)
                style_name = para.style.name if para.style else ""
                text = para.text.strip()
                if not text:
                    continue
                if style_name.startswith('Heading'):
                    if current_heading is not None or current_text:
                        self.sections.append({'type': 'docx_section', 'heading': current_heading, 'text': '\n'.join(current_text)})
                    current_heading = text
                    current_text = []
                else:
                    current_text.append(text)
            elif element.tag.endswith('tbl'):
                table = docx.table.Table(element, self.docx_doc)
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_data:
                        current_text.append(" | ".join(row_data))

        if current_heading is not None or current_text:
            self.sections.append({'type': 'docx_section', 'heading': current_heading, 'text': '\n'.join(current_text)})

        self.total_sections = len(self.sections)
        if self.total_sections == 0:
            full_text = '\n'.join([p.text for p in self.docx_doc.paragraphs if p.text.strip()])
            self.sections.append({'type': 'docx_section', 'heading': _("Full Document"), 'text': full_text})
            self.total_sections = 1

        self.populate_mode_choices_docx()
        self.section_choice.Clear()
        self.section_choice.AppendItems([str(i+1) for i in range(self.total_sections)])
        self.load_section(0)

    def populate_mode_choices_docx(self):
        modes = [_("Plain Text"), _("Formatted")]
        self.mode_choice.SetItems(modes)
        self.mode_choice.SetSelection(0)

    def load_epub(self):
        if ebooklib is None or epub is None or BeautifulSoup is None:
            show_msg(_("ebooklib and BeautifulSoup are required."), _("Error"), True)
            self.Close()
            return
        try:
            self.epub_book = epub.read_epub(self.file_path)
        except Exception as e:
            show_msg(_("Unable to open EPUB: {}").format(str(e)), _("Error"), True)
            self.Close()
            return
        self.sections = []
        for item in self.epub_book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
            content = item.get_body_content().decode('utf-8')
            soup = BeautifulSoup(content, 'html.parser')
            
            for script in soup(["script", "style"]):
                script.extract()

            blocks = []
            for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'li']):
                text = element.get_text(separator=' ', strip=True)
                if text:
                    blocks.append(text)
            
            if not blocks:
                text = soup.get_text(separator='\n', strip=True)
                blocks = [line.strip() for line in text.split('\n') if line.strip()]
                    
            text = '\n\n'.join(blocks)
            text = re.sub(r'\n{3,}', '\n\n', text)
            
            if text.strip():
                self.sections.append({
                    'type': 'epub_section', 
                    'item': item, 
                    'text': text,
                    'id': item.get_id(),
                    'file_name': item.get_name()
                })
                
        self.total_sections = len(self.sections)
        if self.total_sections == 0:
            show_msg(_("EPUB has no readable content."), _("Error"), True)
            self.Close()
            return
        self.populate_mode_choices_epub()
        self.section_choice.Clear()
        self.section_choice.AppendItems([str(i+1) for i in range(self.total_sections)])
        self.load_section(0)

    def populate_mode_choices_epub(self):
        modes = [_("Plain Text"), _("Formatted")]
        self.mode_choice.SetItems(modes)
        self.mode_choice.SetSelection(0)

    def load_section(self, section_num, force=False):
        if section_num < 0 or section_num >= self.total_sections:
            return
        self.current_section = section_num
        text = self.extract_section_text(section_num)
        self.text_ctrl.SetValue(text)
        self.text_ctrl.SetInsertionPoint(0)
        self.section_choice.SetSelection(section_num)
        self.update_toolbar()
        self.update_status()
        self.save_history(section_num)
        if self.split_view_active and hasattr(self, 'split_text'):
            split_section = min(section_num + 1, self.total_sections - 1)
            split_text = self.extract_section_text(split_section)
            self.split_text.SetValue(split_text)

    def extract_section_text(self, section_num):
        if section_num in self.page_cache and not self.force_reload:
            return self.page_cache[section_num]
        profile = self.extraction_profile
        text = ""
        if self.file_type == 'pdf':
            text = self.extract_pdf_text(section_num)
        elif self.file_type == 'docx':
            text = self.extract_docx_text(section_num)
        elif self.file_type == 'epub':
            text = self.extract_epub_text(section_num)
        if self.contrast_mode:
            text = text.upper()
        if self.invert_colors:
            text = text[::-1]
        if profile == "clean":
            text = re.sub(r'\s+', ' ', text).strip()
        elif profile == "simple":
            text = re.sub(r'[^\w\s]', '', text)
        self.page_cache[section_num] = text
        return text

    def extract_pdf_text(self, section_num):
        page = self.pdf_doc[section_num]
        if self.reading_mode == "blocks":
            blocks = page.get_text("blocks")
            lines = [b[4] for b in blocks if b[6] == 0]
            return "\n".join(lines)
        elif self.reading_mode == "words":
            words = page.get_text("words")
            return " ".join([w[4] for w in words])
        elif self.reading_mode == "html":
            return page.get_text("html")
        elif self.reading_mode == "structured":
            blocks = page.get_text("dict")["blocks"]
            structured = []
            for block in blocks:
                if "lines" in block:
                    for line in block["lines"]:
                        spans = " ".join([span["text"] for span in line["spans"]])
                        structured.append(spans)
            return "\n".join(structured)
        else:
            return page.get_text("text")

    def extract_docx_text(self, section_num):
        sec = self.sections[section_num]
        if self.reading_mode == "plain":
            return sec['text']
        else:
            heading = sec.get('heading', '')
            text = sec['text']
            if heading:
                return heading + "\n" + text
            return text

    def extract_epub_text(self, section_num):
        sec = self.sections[section_num]
        if self.reading_mode == "plain":
            return sec['text']
        else:
            return sec['text']

    def update_toolbar(self):
        self.toolbar_text.SetLabel(_("Section {} of {}").format(self.current_section + 1, self.total_sections))

    def update_status(self):
        status = _("Section {cur} of {total} | Mode: {mode} | Zoom: {zoom}%").format(
            cur=self.current_section+1, total=self.total_sections, mode=self.reading_mode, zoom=self.zoom_level
        )
        self.status_bar.SetLabel(status)
        cursor_pos = self.text_ctrl.GetInsertionPoint()
        total_len = len(self.text_ctrl.GetValue())
        self.position_label.SetLabel(_("Pos: {}/{}").format(cursor_pos, total_len))
        ui.message(_("Section {}").format(self.current_section + 1))

    def update_text_display(self):
        font = wx.Font(self.font_size, wx.FONTFAMILY_TELETYPE, wx.FONTSTYLE_NORMAL, wx.FONTWEIGHT_NORMAL)
        self.text_ctrl.SetFont(font)
        if hasattr(self, 'split_text'):
            self.split_text.SetFont(font)

    def on_prev(self, event):
        if self.current_section > 0:
            self.load_section(self.current_section - 1)
        else:
            show_msg(_("Already at the first section"), _("Information"))

    def on_next(self, event):
        if self.current_section < self.total_sections - 1:
            self.load_section(self.current_section + 1)
        else:
            show_msg(_("Already at the last section"), _("Information"))

    def on_section_choice(self, event):
        idx = self.section_choice.GetSelection()
        self.load_section(idx)

    def on_go_to_section(self, event):
        dlg = wx.TextEntryDialog(self, _("Enter section number (1-{}):").format(self.total_sections), _("Go To Section"))
        if dlg.ShowModal() == wx.ID_OK:
            try:
                section = int(dlg.GetValue()) - 1
                if 0 <= section < self.total_sections:
                    self.load_section(section)
                else:
                    show_msg(_("Invalid section number"), _("Error"), True)
            except:
                show_msg(_("Please enter a valid number"), _("Error"), True)
        dlg.Destroy()

    def on_mode_change(self, event):
        mode_map_pdf = {0: "blocks", 1: "text", 2: "words", 3: "html", 4: "structured"}
        mode_map_other = {0: "plain", 1: "formatted"}
        if self.file_type == 'pdf':
            self.reading_mode = mode_map_pdf[self.mode_choice.GetSelection()]
        else:
            self.reading_mode = mode_map_other[self.mode_choice.GetSelection()]
        self.page_cache.clear()
        self.load_section(self.current_section)
        self.save_settings()

    def on_zoom_in(self, event):
        self.zoom_level = min(500, self.zoom_level + 10)
        self.font_size = min(48, self.font_size + 1)
        self.update_text_display()
        self.update_status()
        self.save_settings()
        show_msg(_("Zoom level: {}%").format(self.zoom_level), _("Success"))

    def on_zoom_out(self, event):
        self.zoom_level = max(10, self.zoom_level - 10)
        self.font_size = max(6, self.font_size - 1)
        self.update_text_display()
        self.update_status()
        self.save_settings()
        show_msg(_("Zoom level: {}%").format(self.zoom_level), _("Success"))

    def on_zoom_reset(self, event):
        self.zoom_level = 100
        self.font_size = 12
        self.update_text_display()
        self.update_status()
        self.save_settings()
        show_msg(_("Zoom reset to 100%"), _("Success"))

    def on_high_contrast(self, event):
        self.contrast_mode = not self.contrast_mode
        self.page_cache.clear()
        self.load_section(self.current_section)
        show_msg(_("High contrast {}").format(_("enabled") if self.contrast_mode else _("disabled")), _("Success"))
        self.save_settings()

    def on_invert_colors(self, event):
        self.invert_colors = not self.invert_colors
        self.page_cache.clear()
        self.load_section(self.current_section)
        show_msg(_("Color inversion {}").format(_("enabled") if self.invert_colors else _("disabled")), _("Success"))
        self.save_settings()

    def on_dark_mode(self, event):
        self.dark_mode = not getattr(self, 'dark_mode', False)
        bg = wx.BLACK if self.dark_mode else wx.WHITE
        fg = wx.WHITE if self.dark_mode else wx.BLACK
        self.text_ctrl.SetBackgroundColour(bg)
        self.text_ctrl.SetForegroundColour(fg)
        if hasattr(self, 'split_text'):
            self.split_text.SetBackgroundColour(bg)
            self.split_text.SetForegroundColour(fg)
        self.text_ctrl.Refresh()
        show_msg(_("Dark mode toggled."), _("Success"))

    def on_split_view(self, event):
        self.split_view_active = not self.split_view_active
        if self.split_view_active:
            self.split_text = wx.TextCtrl(self, style=wx.TE_MULTILINE | wx.TE_READONLY | wx.TE_RICH2)
            if self.dark_mode:
                self.split_text.SetBackgroundColour(wx.BLACK)
                self.split_text.SetForegroundColour(wx.WHITE)
            sizer = self.GetSizer()
            sizer.Insert(1, self.split_text, 1, wx.ALL | wx.EXPAND, 5)
            split_section = min(self.current_section + 1, self.total_sections - 1)
            self.split_text.SetValue(self.extract_section_text(split_section))
        else:
            if hasattr(self, 'split_text'):
                self.split_text.Destroy()
                del self.split_text
        self.Layout()
        show_msg(_("Split view {}").format(_("enabled") if self.split_view_active else _("disabled")), _("Success"))

    def on_copy_page(self, event):
        text = self.text_ctrl.GetStringSelection()
        if not text:
            text = self.text_ctrl.GetValue()
        if wx.TheClipboard.Open():
            wx.TheClipboard.SetData(wx.TextDataObject(text))
            wx.TheClipboard.Close()
            show_msg(_("Section text copied to clipboard"), _("Success"))

    def on_copy_all(self, event):
        pd = ProcessingDialog(self, _("Copying all text..."))
        pd.Show()
        def copy_all():
            all_text = []
            for i in range(self.total_sections):
                all_text.append(self.extract_section_text(i))
                wx.CallAfter(pd.update, _("Copying section {}...").format(i+1), int((i+1)/self.total_sections*100))
            full_text = "\n".join(all_text)
            if wx.TheClipboard.Open():
                wx.TheClipboard.SetData(wx.TextDataObject(full_text))
                wx.TheClipboard.Close()
            thread_safe_msg(_("All text copied to clipboard"), _("Success"))
            wx.CallAfter(pd.Destroy)
        threading.Thread(target=copy_all, daemon=True).start()

    def on_search_dialog(self, event):
        search_dlg = SearchDialog(self)
        if search_dlg.ShowModal() == wx.ID_OK:
            query = search_dlg.query
            case = search_dlg.case_sensitive
            whole = search_dlg.whole_words
            regex = search_dlg.regex
            search_range = search_dlg.search_range
            self.perform_search(query, case, whole, regex, search_range)
        search_dlg.Destroy()

    def perform_search(self, query, case, whole, regex, search_range):
        self.search_results = []
        self.current_search_index = -1
        pd = ProcessingDialog(self, _("Searching..."))
        pd.Show()

        data = load_data()
        search_history = data.get("search_history", [])
        if query not in search_history:
            search_history.insert(0, query)
            data["search_history"] = search_history[:20]
            save_data(data)

        def search_thread():
            try:
                flags = 0 if case else re.IGNORECASE
                if not regex:
                    pattern = re.escape(query)
                else:
                    pattern = query
                if whole and not regex:
                    pattern = r"\b" + pattern + r"\b"
                try:
                    regex_obj = re.compile(pattern, flags)
                except:
                    thread_safe_msg(_("Invalid regular expression"), _("Error"), True)
                    return

                results = []
                sections = range(self.total_sections) if search_range == "all" else [self.current_section]

                for section_num in sections:
                    text = self.extract_section_text(section_num)
                    for m in regex_obj.finditer(text):
                        snippet = text[max(0, m.start()-40):m.end()+40]
                        results.append((section_num, snippet.strip(), m.start(), m.group()))
                    wx.CallAfter(pd.update, _("Searching section {}...").format(section_num+1), int((section_num+1)/len(sections)*100))

                self.search_results = results
                if results:
                    thread_safe_msg(_("{} matches found.").format(len(results)), _("Success"))
                    wx.CallAfter(self.show_search_results)
                else:
                    thread_safe_msg(_("No matches found."), _("Information"))
            except Exception as e:
                thread_safe_msg(_("Search error: {}").format(str(e)), _("Error"), True)
            finally:
                wx.CallAfter(pd.Destroy)
        threading.Thread(target=search_thread, daemon=True).start()

    def show_search_results(self):
        dlg = SearchResultsDialog(self, self.search_results)
        if dlg.ShowModal() == wx.ID_OK:
            idx = dlg.result_index
            if idx >= 0:
                self.jump_to_search_result(idx)
        dlg.Destroy()

    def jump_to_search_result(self, index):
        if 0 <= index < len(self.search_results):
            self.current_search_index = index
            section_num, snippet, pos, match_text = self.search_results[index]
            self.load_section(section_num)
            self.text_ctrl.SetInsertionPoint(pos)
            end_pos = pos + len(match_text)
            self.text_ctrl.SetSelection(pos, end_pos)
            self.text_ctrl.ShowPosition(pos)
            show_msg(_("Result {} of {}: {}").format(index+1, len(self.search_results), match_text), _("Information"), silent=True)

    def on_find_next(self, event):
        if self.search_results:
            self.current_search_index = (self.current_search_index + 1) % len(self.search_results)
            self.jump_to_search_result(self.current_search_index)
        else:
            show_msg(_("No search results available"), _("Information"))

    def on_find_previous(self, event):
        if self.search_results:
            self.current_search_index = (self.current_search_index - 1) % len(self.search_results)
            self.jump_to_search_result(self.current_search_index)
        else:
            show_msg(_("No search results available"), _("Information"))

    def on_find_replace(self, event):
        if self.file_type != 'pdf':
            show_msg(_("Find and replace is only available for PDF files."), _("Information"))
            return
        dlg1 = wx.TextEntryDialog(self, _("Find text:"), _("Find and Replace"))
        if dlg1.ShowModal() == wx.ID_OK:
            find_text = dlg1.GetValue()
            dlg2 = wx.TextEntryDialog(self, _("Replace with:"), _("Find and Replace"))
            if dlg2.ShowModal() == wx.ID_OK:
                replace_text = dlg2.GetValue()
                page = self.pdf_doc[self.current_section]
                rects = page.search_for(find_text)
                if rects:
                    for rect in rects:
                        page.add_redact_annot(rect)
                    page.apply_redactions()
                    for rect in rects:
                        page.insert_textbox(rect, replace_text, fontsize=self.font_size, align=0)
                    self.pdf_doc.saveIncr()
                    self.page_cache.clear()
                    self.load_section(self.current_section)
                    show_msg(_("Text replaced on current section."), _("Success"))
                else:
                    show_msg(_("Text not found."), _("Information"))
            dlg2.Destroy()
        dlg1.Destroy()

    def on_add_bookmark(self, event):
        title_dlg = wx.TextEntryDialog(self, _("Bookmark title:"), _("Add Bookmark"))
        if title_dlg.ShowModal() == wx.ID_OK:
            title = title_dlg.GetValue()
            data = load_data()
            data["bookmarks"].append({
                "id": int(time.time() * 1000),
                "file_path": self.file_path,
                "section": self.current_section,
                "title": title,
                "timestamp": time.time()
            })
            save_data(data)
            show_msg(_("Bookmark added."), _("Success"))
        title_dlg.Destroy()

    def on_bookmarks_manager(self, event):
        data = load_data()
        bookmarks = [b for b in data["bookmarks"] if b["file_path"] == self.file_path]
        bookmarks.sort(key=lambda x: x["section"])
        dlg = BookmarksManagerDialog(self, bookmarks)
        if dlg.ShowModal() == wx.ID_OK:
            action = dlg.action
            if action == "jump":
                section = dlg.selected_section
                self.load_section(section)
            elif action == "rename":
                bid = dlg.selected_id
                new_title = dlg.new_title
                for b in data["bookmarks"]:
                    if b["id"] == bid:
                        b["title"] = new_title
                        break
                save_data(data)
                show_msg(_("Bookmark renamed."), _("Success"))
            elif action == "delete":
                bid = dlg.selected_id
                data["bookmarks"] = [b for b in data["bookmarks"] if b["id"] != bid]
                save_data(data)
                show_msg(_("Bookmark deleted."), _("Success"))
        dlg.Destroy()

    def on_export_bookmarks(self, event):
        data = load_data()
        bookmarks = [b for b in data["bookmarks"] if b["file_path"] == self.file_path]
        if not bookmarks:
            show_msg(_("No bookmarks to export."), _("Information"))
            return
        with wx.FileDialog(self, _("Export Bookmarks"), wildcard="Text files (*.txt)|*.txt", style=wx.FD_SAVE | wx.FD_OVERWRITE_PROMPT) as dlg:
            if dlg.ShowModal() == wx.ID_OK:
                with open(dlg.GetPath(), "w", encoding="utf-8") as f:
                    for b in bookmarks:
                        f.write(f"Section {b['section']+1}: {b['title']}\n")
                show_msg(_("Bookmarks exported."), _("Success"))

    def _parse_epub_toc(self, toc_list, level=1):
        parsed = []
        for item in toc_list:
            if isinstance(item, tuple):
                section, sub_items = item
                title = section.title if hasattr(section, 'title') else str(section)
                href = section.href if hasattr(section, 'href') else None
                parsed.append((level, title, href))
                parsed.extend(self._parse_epub_toc(sub_items, level + 1))
            elif hasattr(item, 'title') and hasattr(item, 'href'):
                parsed.append((level, item.title, item.href))
            elif hasattr(item, 'get_name'):
                parsed.append((level, item.get_name(), None))
        return parsed

    def on_toc(self, event):
        if self.file_type == 'pdf':
            toc = self.pdf_doc.get_toc()
            if not toc:
                show_msg(_("No table of contents found."), _("Information"))
                return
            dlg = TocDialog(self, toc, 'pdf')
            if dlg.ShowModal() == wx.ID_OK:
                page = dlg.selected_page
                if page is not None:
                    self.load_section(page - 1)
            dlg.Destroy()
        elif self.file_type == 'docx':
            toc = []
            for i, sec in enumerate(self.sections):
                heading = sec.get('heading', '')
                if heading:
                    level = 1
                    toc.append([level, heading, i+1])
            if not toc:
                show_msg(_("No headings found for TOC."), _("Information"))
                return
            dlg = TocDialog(self, toc, 'docx')
            if dlg.ShowModal() == wx.ID_OK:
                section = dlg.selected_section
                if section is not None:
                    self.load_section(section - 1)
            dlg.Destroy()
        elif self.file_type == 'epub':
            if not self.epub_book:
                return
            parsed_toc = self._parse_epub_toc(self.epub_book.toc)
            toc = []
            for level, title, href in parsed_toc:
                section_idx = 1
                if href:
                    base_href = href.split('#')[0]
                    for i, sec in enumerate(self.sections):
                        if sec.get('file_name') == base_href or sec.get('id') == base_href:
                            section_idx = i + 1
                            break
                toc.append([level, title, section_idx])
            
            if not toc:
                show_msg(_("No table of contents found."), _("Information"))
                return
            dlg = TocDialog(self, toc, 'epub')
            if dlg.ShowModal() == wx.ID_OK:
                section = dlg.selected_section
                if section is not None:
                    self.load_section(section - 1)
            dlg.Destroy()

    def on_ocr(self, event):
        if self.file_type != 'pdf':
            show_msg(_("OCR is only available for PDF files."), _("Information"))
            return
        dlg = OCRDialog(self)
        if dlg.ShowModal() == wx.ID_OK:
            pages = dlg.get_pages()
            self.run_ocr(pages)
        dlg.Destroy()

    def run_ocr(self, pages):
        if pytesseract is None or Image is None:
            show_msg(_("pytesseract and Pillow are required for OCR."), _("Error"), True)
            return
        if not pages:
            return

        tesseract_found = False
        possible_paths = [
            r"C:\Program Files\Tesseract-OCR\tesseract.exe",
            r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
        ]
        for p in possible_paths:
            if os.path.isfile(p):
                pytesseract.pytesseract.tesseract_cmd = p
                tesseract_found = True
                break

        if not tesseract_found:
            try:
                pytesseract.get_tesseract_version()
                tesseract_found = True
            except:
                self.handle_missing_tesseract()
                return

        pd = ProcessingDialog(self, _("OCR in progress..."))
        pd.Show()
        def ocr_thread():
            total = len(pages)
            for i, page_num in enumerate(pages):
                if page_num in self.ocr_cache:
                    continue
                try:
                    pix = self.pdf_doc[page_num].get_pixmap(dpi=300)
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    img = img.convert('L')
                    img = img.filter(ImageFilter.SHARPEN)
                    enhancer = ImageEnhance.Contrast(img)
                    img = enhancer.enhance(2.0)
                    text = pytesseract.image_to_string(img, lang='eng+fra+deu+spa')
                    self.ocr_cache[page_num] = text
                    wx.CallAfter(pd.update, _("OCR processed page {}").format(page_num+1), int((i+1)/total*100))
                except Exception as e:
                    thread_safe_msg(_("OCR failed on page {}: {}").format(page_num+1, str(e)), _("Error"), True)
            thread_safe_msg(_("OCR completed."), _("Success"))
            wx.CallAfter(setattr, self, 'extraction_profile', 'ocr_only')
            wx.CallAfter(self.page_cache.clear)
            wx.CallAfter(self.load_section, self.current_section)
            wx.CallAfter(pd.Destroy)
        threading.Thread(target=ocr_thread, daemon=True).start()

    def handle_missing_tesseract(self):
        dlg = wx.MessageDialog(self,
            _("The Tesseract installation could not be found in the default path. Would you like to browse yourself or download the installer?"),
            _("Missing Tesseract"),
            wx.YES_NO | wx.CANCEL | wx.ICON_QUESTION
        )
        if hasattr(dlg, 'SetYesNoCancelLabels'):
            dlg.SetYesNoCancelLabels(_("Browse Tesseract executable"), _("Download Tesseract installer"), _("Not now"))
        else:
            dlg.SetYesNoLabels(_("Browse"), _("Download"))
            
        result = dlg.ShowModal()
        dlg.Destroy()
        if result == wx.ID_YES:
            with wx.FileDialog(self, _("Locate tesseract.exe"), wildcard="Executable (*.exe)|*.exe") as file_dlg:
                if file_dlg.ShowModal() == wx.ID_OK:
                    pytesseract.pytesseract.tesseract_cmd = file_dlg.GetPath()
                    show_msg(_("Tesseract path set."), _("Success"))
        elif result == wx.ID_NO:
            self.download_tesseract()

    def download_tesseract(self):
        url = "https://github.com/tesseract-ocr/tesseract/releases/download/5.5.0/tesseract-ocr-w64-setup-5.5.0.20241111.exe"
        downloads_dir = os.path.join(os.path.expanduser("~"), "Downloads")
        if not os.path.exists(downloads_dir):
            os.makedirs(downloads_dir)
        dest_path = os.path.join(downloads_dir, "tesseract-ocr-w64-setup-5.5.0.20241111.exe")
        pd = ProcessingDialog(self, _("Downloading..."))
        pd.Show()
        def download():
            try:
                def reporthook(blocknum, blocksize, totalsize):
                    readsofar = blocknum * blocksize
                    if totalsize > 0:
                        percent = min(int(readsofar * 100 / totalsize), 100)
                        wx.CallAfter(pd.update, _("Downloading... {}%").format(percent), percent)
                urllib.request.urlretrieve(url, dest_path, reporthook)
                wx.CallAfter(pd.Destroy)
                wx.CallAfter(lambda: subprocess.Popen([dest_path], shell=True))
                thread_safe_msg(_("Download completed. Installer launched."), _("Success"))
            except Exception as e:
                wx.CallAfter(pd.Destroy)
                thread_safe_msg(_("Download failed: {}").format(str(e)), _("Error"), True)
        threading.Thread(target=download, daemon=True).start()

    def on_export_menu(self, event):
        export_dlg = ExportDialog(self)
        if export_dlg.ShowModal() == wx.ID_OK:
            export_type = export_dlg.export_type
            scope = export_dlg.scope
            include_notes = export_dlg.include_notes
            include_ocr = export_dlg.include_ocr
            self.do_export(export_type, scope, include_notes, include_ocr)
        export_dlg.Destroy()

    def do_export(self, export_type, scope, include_notes=False, include_ocr=False):
        if scope == "current":
            sections = [self.current_section]
        elif scope == "all":
            sections = list(range(self.total_sections))
        else:
            data = load_data()
            tags = data.get("doc_tags", {}).get(self.file_path, {})
            sections = list(tags.keys())

        if export_type == "txt":
            self.export_txt(sections, include_notes, include_ocr)
        elif export_type == "docx":
            self.export_docx(sections, include_notes, include_ocr)
        elif export_type == "audiobook":
            self.generate_audiobook(sections)
        elif export_type == "csv":
            self.export_csv(sections)
        elif export_type == "json":
            self.export_json(sections, include_notes)
        elif export_type == "html":
            self.export_html(sections, include_notes)

    def export_txt(self, sections, include_notes=False, include_ocr=False):
        default_dir = self.settings.get("export_dir", "")
        with wx.FileDialog(self, _("Save as text"), defaultDir=default_dir, wildcard="Text files (*.txt)|*.txt", style=wx.FD_SAVE | wx.FD_OVERWRITE_PROMPT) as dlg:
            if dlg.ShowModal() == wx.ID_OK:
                path = dlg.GetPath()
                pd = ProcessingDialog(self, _("Exporting text..."))
                pd.Show()
                def write():
                    try:
                        with open(path, "w", encoding="utf-8") as f:
                            for i, s in enumerate(sections):
                                text = self.extract_section_text(s)
                                if include_ocr and s in self.ocr_cache:
                                    text += "\n[OCR]\n" + self.ocr_cache[s]
                                f.write(f"=== Section {s+1} ===\n{text}\n\n")
                                if include_notes:
                                    data = load_data()
                                    notes = [n for n in data["notes"] if n["file_path"] == self.file_path and n["section"] == s]
                                    if notes:
                                        f.write("--- Notes ---\n")
                                        for note in notes:
                                            f.write(f"  * {note['note_text']}\n")
                                        f.write("\n")
                                wx.CallAfter(pd.update, _("Exporting section {}...").format(s+1), int((i+1)/len(sections)*100))
                        thread_safe_msg(_("Text exported successfully."), _("Success"))
                    except Exception as e:
                        thread_safe_msg(_("Export error: {}").format(str(e)), _("Error"), True)
                    finally:
                        wx.CallAfter(pd.Destroy)
                threading.Thread(target=write, daemon=True).start()

    def export_docx(self, sections, include_notes=False, include_ocr=False):
        if docx is None:
            show_msg(_("python-docx is required for DOCX export."), _("Error"), True)
            return
        default_dir = self.settings.get("export_dir", "")
        with wx.FileDialog(self, _("Save as DOCX"), defaultDir=default_dir, wildcard="Word files (*.docx)|*.docx", style=wx.FD_SAVE | wx.FD_OVERWRITE_PROMPT) as dlg:
            if dlg.ShowModal() == wx.ID_OK:
                path = dlg.GetPath()
                pd = ProcessingDialog(self, _("Exporting to DOCX..."))
                pd.Show()
                def write():
                    try:
                        document = docx.Document()
                        for i, s in enumerate(sections):
                            document.add_heading(f"Section {s+1}", level=1)
                            text = self.extract_section_text(s)
                            document.add_paragraph(text)
                            if include_ocr and s in self.ocr_cache:
                                document.add_heading("OCR Text", level=2)
                                document.add_paragraph(self.ocr_cache[s])
                            if include_notes:
                                data = load_data()
                                notes = [n for n in data["notes"] if n["file_path"] == self.file_path and n["section"] == s]
                                if notes:
                                    document.add_heading("Notes", level=2)
                                    for note in notes:
                                        document.add_paragraph(note["note_text"], style="List Bullet")
                            wx.CallAfter(pd.update, _("Exporting section {}...").format(s+1), int((i+1)/len(sections)*100))
                        document.save(path)
                        thread_safe_msg(_("DOCX exported successfully."), _("Success"))
                    except Exception as e:
                        thread_safe_msg(_("Export error: {}").format(str(e)), _("Error"), True)
                    finally:
                        wx.CallAfter(pd.Destroy)
                threading.Thread(target=write, daemon=True).start()

    def export_csv(self, sections):
        if self.file_type != 'pdf' or pdfplumber is None:
            show_msg(_("CSV export is only available for PDF files with tables."), _("Information"))
            return
        default_dir = self.settings.get("export_dir", "")
        with wx.FileDialog(self, _("Save as CSV"), defaultDir=default_dir, wildcard="CSV files (*.csv)|*.csv", style=wx.FD_SAVE | wx.FD_OVERWRITE_PROMPT) as dlg:
            if dlg.ShowModal() == wx.ID_OK:
                path = dlg.GetPath()
                pd = ProcessingDialog(self, _("Exporting tables to CSV..."))
                pd.Show()
                def extract():
                    try:
                        import pandas
                    except ImportError:
                        thread_safe_msg(_("pandas is required for CSV export."), _("Error"), True)
                        wx.CallAfter(pd.Destroy)
                        return
                    all_tables = []
                    for i, p in enumerate(sections):
                        with pdfplumber.open(self.file_path) as pdf:
                            page = pdf.pages[p]
                            tables = page.extract_tables()
                            for table in tables:
                                all_tables.append(table)
                        wx.CallAfter(pd.update, _("Extracting tables from section {}...").format(p+1), int((i+1)/len(sections)*100))
                    if all_tables:
                        combined = []
                        for table in all_tables:
                            combined.extend(table)
                        df = pandas.DataFrame(combined)
                        df.to_csv(path, index=False)
                        thread_safe_msg(_("CSV exported."), _("Success"))
                    else:
                        thread_safe_msg(_("No tables found."), _("Information"))
                    wx.CallAfter(pd.Destroy)
                threading.Thread(target=extract, daemon=True).start()

    def export_json(self, sections, include_notes=False):
        default_dir = self.settings.get("export_dir", "")
        with wx.FileDialog(self, _("Save as JSON"), defaultDir=default_dir, wildcard="JSON files (*.json)|*.json", style=wx.FD_SAVE | wx.FD_OVERWRITE_PROMPT) as dlg:
            if dlg.ShowModal() == wx.ID_OK:
                path = dlg.GetPath()
                pd = ProcessingDialog(self, _("Exporting to JSON..."))
                pd.Show()
                def write():
                    try:
                        export_data = {
                            "file_path": self.file_path,
                            "total_sections": self.total_sections,
                            "exported_sections": [],
                            "metadata": self.pdf_doc.metadata if self.file_type == 'pdf' else {}
                        }
                        for i, s in enumerate(sections):
                            section_data = {
                                "section_number": s + 1,
                                "text": self.extract_section_text(s),
                                "ocr_text": self.ocr_cache.get(s, ""),
                                "notes": []
                            }
                            if include_notes:
                                data = load_data()
                                section_data["notes"] = [n["note_text"] for n in data["notes"] if n["file_path"] == self.file_path and n["section"] == s]
                            export_data["exported_sections"].append(section_data)
                            wx.CallAfter(pd.update, _("Exporting section {}...").format(s+1), int((i+1)/len(sections)*100))
                        with open(path, "w", encoding="utf-8") as f:
                            json.dump(export_data, f, ensure_ascii=False, indent=2)
                        thread_safe_msg(_("JSON exported successfully."), _("Success"))
                    except Exception as e:
                        thread_safe_msg(_("Export error: {}").format(str(e)), _("Error"), True)
                    finally:
                        wx.CallAfter(pd.Destroy)
                threading.Thread(target=write, daemon=True).start()

    def export_html(self, sections, include_notes=False):
        default_dir = self.settings.get("export_dir", "")
        with wx.FileDialog(self, _("Save as HTML"), defaultDir=default_dir, wildcard="HTML files (*.html)|*.html", style=wx.FD_SAVE | wx.FD_OVERWRITE_PROMPT) as dlg:
            if dlg.ShowModal() == wx.ID_OK:
                path = dlg.GetPath()
                pd = ProcessingDialog(self, _("Exporting to HTML..."))
                pd.Show()
                def write():
                    try:
                        html_content = ["<html><head><title>Document Export</title></head><body>"]
                        for i, s in enumerate(sections):
                            html_content.append(f"<h1>Section {s+1}</h1>")
                            text = self.extract_section_text(s).replace("\n", "<br>")
                            html_content.append(f"<p>{text}</p>")
                            if include_notes:
                                data = load_data()
                                notes = [n for n in data["notes"] if n["file_path"] == self.file_path and n["section"] == s]
                                if notes:
                                    html_content.append("<h2>Notes</h2>")
                                    for note in notes:
                                        html_content.append(f"<p>{note['note_text']}</p>")
                            wx.CallAfter(pd.update, _("Exporting section {}...").format(s+1), int((i+1)/len(sections)*100))
                        html_content.append("</body></html>")
                        with open(path, "w", encoding="utf-8") as f:
                            f.write("\n".join(html_content))
                        thread_safe_msg(_("HTML exported successfully."), _("Success"))
                    except Exception as e:
                        thread_safe_msg(_("Export error: {}").format(str(e)), _("Error"), True)
                    finally:
                        wx.CallAfter(pd.Destroy)
                threading.Thread(target=write, daemon=True).start()

    def generate_audiobook(self, sections):
        if edge_tts is None:
            show_msg(_("edge_tts library not installed or missing dependencies."), _("Error"), True)
            return
        dlg = wx.TextEntryDialog(self, _("Enter section numbers to include, separated by spaces (e.g., '1 2 4'). Leave blank for selected sections:"), _("Audiobook Sections"))
        if dlg.ShowModal() == wx.ID_OK:
            input_str = dlg.GetValue().strip()
            if input_str:
                try:
                    selected = [int(x.strip())-1 for x in input_str.split()]
                    invalid = [x for x in selected if x < 0 or x >= self.total_sections]
                    if invalid:
                        show_msg(_("Invalid section numbers: {}").format(", ".join(str(x+1) for x in invalid)), _("Error"), True)
                        dlg.Destroy()
                        return
                    sections = selected
                except:
                    show_msg(_("Invalid input. Use numbers separated by spaces."), _("Error"), True)
                    dlg.Destroy()
                    return
        dlg.Destroy()
        if not sections:
            show_msg(_("No sections selected."), _("Error"), True)
            return
        voice = self.settings.get("voice", "en-US-AriaNeural")
        speed = self.settings.get("tts_speed", "+0%")
        pitch = self.settings.get("tts_pitch", "+0Hz")
        vol_val = self.settings.get("tts_volume", 100) - 100
        vol_str = f"+{vol_val}%" if vol_val >= 0 else f"{vol_val}%"
        use_chunking = self.settings.get("tts_chunking", True)
        default_dir = self.settings.get("export_dir", "")
        with wx.FileDialog(self, _("Save audiobook"), defaultDir=default_dir, wildcard="MP3 files (*.mp3)|*.mp3", style=wx.FD_SAVE | wx.FD_OVERWRITE_PROMPT) as dlg:
            if dlg.ShowModal() == wx.ID_OK:
                path = dlg.GetPath()
                pd = ProcessingDialog(self, _("Generating audiobook..."))
                pd.Show()
                def generate():
                    try:
                        async def build_audiobook():
                            with open(path, 'wb') as f:
                                for i, s in enumerate(sections):
                                    text = self.extract_section_text(s)
                                    full_text = f"Page number {s+1}. {text}"
                                    if not full_text.strip(): 
                                        continue
                                    if use_chunking:
                                        chunks = [full_text[j:j+4000] for j in range(0, len(full_text), 4000)]
                                        for chunk in chunks:
                                            communicate = edge_tts.Communicate(chunk, voice, rate=speed, pitch=pitch, volume=vol_str)
                                            async for audio_chunk in communicate.stream():
                                                if audio_chunk["type"] == "audio":
                                                    f.write(audio_chunk["data"])
                                    else:
                                        communicate = edge_tts.Communicate(full_text, voice, rate=speed, pitch=pitch, volume=vol_str)
                                        async for audio_chunk in communicate.stream():
                                            if audio_chunk["type"] == "audio":
                                                f.write(audio_chunk["data"])
                                    wx.CallAfter(pd.update, _("Generating section {}...").format(s+1), int((i+1)/len(sections)*100))
                        asyncio.run(build_audiobook())
                        wx.CallAfter(pd.Destroy)
                        thread_safe_msg(_("Audiobook generated successfully."), _("Success"))
                    except Exception as e:
                        wx.CallAfter(pd.Destroy)
                        thread_safe_msg(_("Generation error: {}").format(str(e)), _("Error"), True)
                threading.Thread(target=generate, daemon=True).start()

    def on_add_note(self, event):
        dlg = wx.TextEntryDialog(self, _("Note text:"), _("Add Note"), style=wx.TE_MULTILINE)
        if dlg.ShowModal() == wx.ID_OK:
            note = dlg.GetValue()
            data = load_data()
            data["notes"].append({
                "id": int(time.time() * 1000),
                "file_path": self.file_path,
                "section": self.current_section,
                "note_text": note,
                "timestamp": time.time()
            })
            save_data(data)
            show_msg(_("Note added."), _("Success"))
        dlg.Destroy()

    def on_annotations(self, event):
        if self.file_type != 'pdf':
            show_msg(_("Annotations are only available for PDF files."), _("Information"))
            return
        data = load_data()
        annotations = [a for a in data.get("annotations", []) if a["file_path"] == self.file_path]
        dlg = AnnotationsDialog(self, annotations)
        dlg.ShowModal()
        dlg.Destroy()

    def on_highlights(self, event):
        data = load_data()
        highlights = [h for h in data.get("highlights", []) if h["file_path"] == self.file_path]
        dlg = HighlightsDialog(self, highlights, self.pdf_doc if self.file_type == 'pdf' else None, self.file_path)
        if dlg.ShowModal() == wx.ID_OK:
            self.page_cache.clear()
            self.load_section(self.current_section)
        dlg.Destroy()

    def on_highlight_selection(self, event):
        selected_text = self.text_ctrl.GetStringSelection()
        if not selected_text:
            show_msg(_("No text selected."), _("Information"))
            return
        section_num = self.current_section
        data = load_data()
        highlight_id = int(time.time() * 1000)
        data["highlights"].append({
            "id": highlight_id,
            "file_path": self.file_path,
            "section": section_num,
            "text": selected_text,
            "timestamp": time.time()
        })
        save_data(data)
        if self.file_type == 'pdf':
            page = self.pdf_doc[section_num]
            rects = page.search_for(selected_text)
            for rect in rects:
                annot = page.add_highlight_annot(rect)
                annot.set_info(title="NVDA Highlight")
                annot.update()
            self.pdf_doc.saveIncr()
        show_msg(_("Highlight added."), _("Success"))

    def on_read_aloud(self, event):
        if self.read_aloud:
            self.stop_read_aloud()
        else:
            self.start_read_aloud()

    def start_read_aloud(self):
        if edge_tts is None:
            show_msg(_("edge_tts library not installed or missing dependencies."), _("Error"), True)
            return

        self.stop_read_aloud()

        self.read_aloud = True
        self.read_aloud_btn.SetLabel(_("&Stop"))

        text = self.extract_section_text(self.current_section)

        if not text.strip():
            show_msg(_("No text to read."), _("Error"), True)
            self.stop_read_aloud()
            return

        pd = ProcessingDialog(self, _("Converting text to speech..."))
        pd.Show()

        def tts_thread():
            try:
                voice = self.settings.get("voice", "en-US-AriaNeural")
                speed = self.settings.get("tts_speed", "+0%")
                pitch = self.settings.get("tts_pitch", "+0Hz")

                vol_val = self.settings.get("tts_volume", 100) - 100
                vol_str = f"+{vol_val}%" if vol_val >= 0 else f"{vol_val}%"

                temp_file = os.path.join(
                    tempfile.gettempdir(),
                    f"nvda_doc_reader_{int(time.time())}.mp3"
                )

                async def generate():
                    communicate = edge_tts.Communicate(
                        text,
                        voice,
                        rate=speed,
                        pitch=pitch,
                        volume=vol_str
                    )
                    await communicate.save(temp_file)

                asyncio.run(generate())

                if not self.read_aloud:
                    try:
                        os.remove(temp_file)
                    except:
                        pass
                    return

                self.current_audio_file = temp_file

                wx.CallAfter(pd.Destroy)
                wx.CallAfter(self.play_audio, temp_file)

            except Exception as e:
                wx.CallAfter(pd.Destroy)
                thread_safe_msg(_("TTS failed: {}").format(str(e)), _("Error"), True)
                wx.CallAfter(self.stop_read_aloud)

        threading.Thread(target=tts_thread, daemon=True).start()

    def play_audio(self, audio_file):
        if not self.read_aloud:
            return

        ffplay_path = os.path.join(base_path, "libs", "ffplay.exe")

        if not os.path.isfile(ffplay_path):
            show_msg(_("ffplay.exe not found."), _("Error"), True)
            self.stop_read_aloud()
            return

        try:
            self.ffplay_process = subprocess.Popen(
                [
                    ffplay_path,
                    "-nodisp",
                    "-autoexit",
                    "-loglevel",
                    "quiet",
                    audio_file
                ],
                stdin=subprocess.DEVNULL,
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
                creationflags=subprocess.CREATE_NO_WINDOW
            )

            threading.Thread(
                target=self._wait_for_ffplay,
                daemon=True
            ).start()

            ui.message(_("Reading aloud"))

        except Exception as e:
            show_msg(_("Playback failed: {}").format(str(e)), _("Error"), True)
            self.stop_read_aloud()

    def _wait_for_ffplay(self):
        try:
            process = self.ffplay_process

            if process is not None:
                process.wait()

            if self.read_aloud:
                wx.CallAfter(self.stop_read_aloud)

        finally:
            audio_file = self.current_audio_file

            self.ffplay_process = None
            self.current_audio_file = None

            if audio_file and os.path.exists(audio_file):
                try:
                    os.remove(audio_file)
                except:
                    pass

    def stop_read_aloud(self):
        self.read_aloud = False

        if self.read_aloud_btn:
            self.read_aloud_btn.SetLabel(_("&Read"))

        if self.ffplay_process is not None:
            try:
                self.ffplay_process.terminate()
                self.ffplay_process.wait(timeout=2)
            except:
                try:
                    self.ffplay_process.kill()
                except:
                    pass

        self.ffplay_process = None

        audio_file = self.current_audio_file
        self.current_audio_file = None

        if audio_file and os.path.exists(audio_file):
            try:
                os.remove(audio_file)
            except:
                pass

        ui.message(_("Reading aloud stopped"))

    def on_auto_scroll(self, event):
        if self.auto_scroll:
            self.stop_auto_scroll()
        else:
            self.start_auto_scroll()

    def start_auto_scroll(self):
        self.auto_scroll = True
        speed = max(1, 10 - self.reading_speed)
        self.auto_scroll_timer = wx.Timer(self)
        self.Bind(wx.EVT_TIMER, self.on_auto_scroll_tick, self.auto_scroll_timer)
        self.auto_scroll_timer.Start(speed * 1000)
        show_msg(_("Auto-scroll started"), _("Information"), silent=True)

    def stop_auto_scroll(self):
        self.auto_scroll = False
        if self.auto_scroll_timer:
            self.auto_scroll_timer.Stop()
            self.auto_scroll_timer = None
        show_msg(_("Auto-scroll stopped"), _("Information"), silent=True)

    def on_auto_scroll_tick(self, event):
        if self.current_section < self.total_sections - 1:
            self.load_section(self.current_section + 1)
        else:
            self.stop_auto_scroll()
            show_msg(_("End of document reached"), _("Information"))

    def on_extraction_profile(self, event):
        dlg = ExtractionProfileDialog(self)
        if dlg.ShowModal() == wx.ID_OK:
            self.extraction_profile = dlg.selected_profile
            self.page_cache.clear()
            self.load_section(self.current_section)
            self.save_settings()
            show_msg(_("Extraction profile changed."), _("Success"))
        dlg.Destroy()

    def on_compare_sections(self, event):
        dlg = wx.TextEntryDialog(self, _("Enter section numbers to compare (e.g., '1,5' or '1-3'):"), _("Compare Sections"))
        if dlg.ShowModal() == wx.ID_OK:
            sections = self.parse_section_range(dlg.GetValue())
            if len(sections) >= 2:
                self.show_section_comparison(sections)
            else:
                show_msg(_("Please specify at least 2 sections"), _("Error"), True)
        dlg.Destroy()

    def parse_section_range(self, range_str):
        sections = set()
        parts = range_str.split(',')
        for part in parts:
            part = part.strip()
            if '-' in part:
                try:
                    start, end = part.split('-')
                    start = int(start.strip()) - 1
                    end = int(end.strip()) - 1
                    sections.update(range(max(0, start), min(self.total_sections, end + 1)))
                except:
                    pass
            else:
                try:
                    section = int(part) - 1
                    if 0 <= section < self.total_sections:
                        sections.add(section)
                except:
                    pass
        return sorted(list(sections))

    def show_section_comparison(self, sections):
        comparison = []
        for s in sections:
            text = self.extract_section_text(s)
            word_count = len(text.split())
            char_count = len(text)
            comparison.append(_("Section {}: {} words, {} characters").format(s+1, word_count, char_count))
        dlg = wx.MessageDialog(self, "\n".join(comparison), _("Section Comparison"), wx.OK)
        dlg.ShowModal()
        dlg.Destroy()

    def on_statistics(self, event):
        pd = ProcessingDialog(self, _("Calculating statistics..."))
        pd.Show()
        def calc():
            try:
                total_words = 0
                total_chars = 0
                total_images = 0
                total_links = 0
                for i in range(self.total_sections):
                    text = self.extract_section_text(i)
                    total_words += len(text.split())
                    total_chars += len(text)
                    if self.file_type == 'pdf':
                        total_images += len(self.pdf_doc[i].get_images())
                        total_links += len(self.pdf_doc[i].get_links())
                    wx.CallAfter(pd.update, _("Analyzing section {}...").format(i+1), int((i+1)/self.total_sections*100))
                stats = _(
                    "Document Statistics:\n"
                    "Total Sections: {}\n"
                    "Total Words: {}\n"
                    "Total Characters: {}\n"
                    "Total Images: {}\n"
                    "Total Links: {}\n"
                    "Average Words/Section: {:.1f}\n"
                    "Average Chars/Section: {:.1f}"
                ).format(self.total_sections, total_words, total_chars, total_images, total_links,
                        total_words/max(1, self.total_sections), total_chars/max(1, self.total_sections))
                wx.CallAfter(self.show_statistics_dialog, stats)
            except Exception as e:
                thread_safe_msg(_("Statistics error: {}").format(str(e)), _("Error"), True)
            finally:
                wx.CallAfter(pd.Destroy)
        threading.Thread(target=calc, daemon=True).start()

    def show_statistics_dialog(self, stats):
        dlg = wx.MessageDialog(self, stats, _("Document Statistics"), wx.OK)
        dlg.ShowModal()
        dlg.Destroy()

    def on_translate_section(self, event):
        if GoogleTranslator is None:
            show_msg(_("deep_translator library not installed. Install deep-translator in libs folder."), _("Error"), True)
            return

        section_num = self.current_section
        if self.translated_sections.get(section_num, False):
            self.reverse_translation(section_num)
        else:
            lang_names = [lang["name"] for lang in languages.LANGUAGES]
            dlg = wx.SingleChoiceDialog(self, _("Select a target language to translate current section content"),
                                        _("Translate Section"), lang_names)
            if dlg.ShowModal() == wx.ID_OK:
                selected_idx = dlg.GetSelection()
                target_code = languages.LANGUAGES[selected_idx]["code"]
                dlg.Destroy()
                self.perform_translation(section_num, target_code)
            else:
                dlg.Destroy()

    def perform_translation(self, section_num, target_code):
        original_text = self.text_ctrl.GetValue()
        if not original_text.strip():
            show_msg(_("No text to translate."), _("Information"))
            return

        self.original_texts[section_num] = original_text

        def split_text_by_words(text, max_len=5000):
            words = text.split()
            chunks = []
            current_chunk = []
            current_len = 0
            for word in words:
                word_len = len(word)
                if current_len + word_len + (1 if current_chunk else 0) > max_len and current_chunk:
                    chunks.append(" ".join(current_chunk))
                    current_chunk = [word]
                    current_len = word_len
                else:
                    current_chunk.append(word)
                    current_len += word_len + (1 if current_chunk else 0)
            if current_chunk:
                chunks.append(" ".join(current_chunk))
            return chunks

        chunks = split_text_by_words(original_text)
        if not chunks:
            show_msg(_("No text to translate."), _("Information"))
            return

        pd = ProcessingDialog(self, _("Translating current section..."))
        pd.Show()

        def translate_thread():
            try:
                translated_parts = []
                total = len(chunks)
                for i, chunk in enumerate(chunks):
                    translator = GoogleTranslator(source='auto', target=target_code)
                    translated = translator.translate(chunk)
                    translated_parts.append(translated)
                    wx.CallAfter(pd.update, _("Translating chunk {} of {}...").format(i+1, total),
                                 int((i+1)/total*100))
                full_translated = " ".join(translated_parts)

                self.page_cache[section_num] = full_translated
                self.translated_sections[section_num] = True

                wx.CallAfter(self.text_ctrl.SetValue, full_translated)
                wx.CallAfter(self.update_status)
                wx.CallAfter(pd.Destroy)
                thread_safe_msg(_("Section translated successfully."), _("Success"))
            except Exception as e:
                wx.CallAfter(pd.Destroy)
                thread_safe_msg(_("Translation error: {}").format(str(e)), _("Error"), True)
                if section_num in self.original_texts:
                    del self.original_texts[section_num]

        threading.Thread(target=translate_thread, daemon=True).start()

    def reverse_translation(self, section_num):
        if section_num not in self.original_texts:
            show_msg(_("Original text not available."), _("Error"), True)
            return
        original_text = self.original_texts.pop(section_num)
        self.page_cache[section_num] = original_text
        self.translated_sections.pop(section_num, None)
        self.text_ctrl.SetValue(original_text)
        self.update_status()
        show_msg(_("Reversed to original language."), _("Success"))

    def on_extract_links(self, event):
        if self.file_type != 'pdf':
            show_msg(_("Link extraction is only available for PDF files."), _("Information"))
            return
        links = []
        for i in range(self.total_sections):
            for link in self.pdf_doc[i].get_links():
                if "uri" in link:
                    links.append(f"Section {i+1}: {link['uri']}")
        if not links:
            show_msg(_("No links found."), _("Information"))
            return
        default_dir = self.settings.get("export_dir", "")
        with wx.FileDialog(self, _("Save Links"), defaultDir=default_dir, wildcard="Text files (*.txt)|*.txt", style=wx.FD_SAVE | wx.FD_OVERWRITE_PROMPT) as dlg:
            if dlg.ShowModal() == wx.ID_OK:
                with open(dlg.GetPath(), "w", encoding="utf-8") as f:
                    f.write("\n".join(links))
                show_msg(_("Links extracted successfully."), _("Success"))

    def on_voice_preview(self, event):
        if edge_tts is None:
            show_msg(_("edge_tts library not installed or missing dependencies."), _("Error"), True)
            return
        voice = self.settings.get("voice", "en-US-AriaNeural")
        speed = self.settings.get("tts_speed", "+0%")
        pitch = self.settings.get("tts_pitch", "+0Hz")
        vol_val = self.settings.get("tts_volume", 100) - 100
        vol_str = f"+{vol_val}%" if vol_val >= 0 else f"{vol_val}%"
        text = _("Hello, This is all your's. I am ready to help you read your documents.")
        pd = ProcessingDialog(self, _("Generating preview..."))
        pd.Show()
        def preview():
            try:
                async def gen():
                    temp_file = os.path.join(tempfile.gettempdir(), f"nvda_voice_preview_{int(time.time())}.mp3")
                    communicate = edge_tts.Communicate(text, voice, rate=speed, pitch=pitch, volume=vol_str)
                    await communicate.save(temp_file)
                    self.current_audio_file = temp_file
                    wx.CallAfter(pd.Destroy)
                    self.read_aloud = True
                    wx.CallAfter(self.play_audio, temp_file)
                asyncio.run(gen())
            except Exception as e:
                wx.CallAfter(pd.Destroy)
                thread_safe_msg(_("Preview error: {}").format(str(e)), _("Error"), True)
        threading.Thread(target=preview, daemon=True).start()

    def on_import(self, event):
        wildcard = "Supported files (*.pdf;*.docx;*.epub)|*.pdf;*.docx;*.epub|PDF files (*.pdf)|*.pdf|Word files (*.docx)|*.docx|EPUB files (*.epub)|*.epub"
        with wx.FileDialog(self, _("Choose Document"), wildcard=wildcard) as file_dialog:
            if file_dialog.ShowModal() == wx.ID_OK:
                new_path = file_dialog.GetPath()
                self.file_path = new_path
                if self.pdf_doc:
                    self.pdf_doc.close()
                self.pdf_doc = None
                self.docx_doc = None
                self.epub_book = None
                self.page_cache.clear()
                self.ocr_cache.clear()
                self.search_results.clear()
                self.translated_sections.clear()
                self.original_texts.clear()
                self.sections.clear()
                self.load_document()
                show_msg(_("Opened: {}").format(os.path.basename(new_path)), _("Success"))

    def on_properties(self, event):
        if self.file_type == 'pdf':
            dlg = MetadataDialog(self, self.file_path, self.pdf_doc)
        else:
            show_msg(_("Properties are only available for PDF files."), _("Information"))
            return
        dlg.ShowModal()
        dlg.Destroy()

    def on_protect_pdf(self, event):
        if self.file_type != 'pdf':
            show_msg(_("This feature is only available for PDF files."), _("Information"))
            return
        dlg = ProtectPdfDialog(self)
        if dlg.ShowModal() == wx.ID_OK:
            pw = dlg.get_password()
            dlg.Destroy()
            default_dir = self.settings.get("export_dir", "")
            with wx.FileDialog(self, _("Save Protected PDF"), defaultDir=default_dir, wildcard="PDF files (*.pdf)|*.pdf", style=wx.FD_SAVE | wx.FD_OVERWRITE_PROMPT) as fdlg:
                if fdlg.ShowModal() == wx.ID_OK:
                    self.pdf_doc.save(fdlg.GetPath(), encryption=fitz.PDF_ENCRYPT_AES_256, owner_pw=pw, user_pw=pw)
                    show_msg(_("PDF protected successfully."), _("Success"))
        else:
            dlg.Destroy()

    def on_help(self, event):
        dlg = HelpDialog(self)
        dlg.ShowModal()
        dlg.Destroy()

    def on_about(self, event):
        dlg = AboutDialog(self)
        dlg.ShowModal()
        dlg.Destroy()

    def on_add_annotation(self, event):
        if self.file_type != 'pdf':
            show_msg(_("Annotations are only available for PDF files."), _("Information"))
            return
        dlg = AnnotationEntryDialog(self)
        if dlg.ShowModal() == wx.ID_OK:
            ann_text = dlg.annotation_text
            dlg.Destroy()
            try:
                page = self.pdf_doc[self.current_section]
                annot = page.add_text_annot((100, 100), ann_text)
                annot.update()
                default_dir = self.settings.get("export_dir", "")
                with wx.FileDialog(self, _("Save annotated PDF"), defaultDir=default_dir, wildcard="PDF files (*.pdf)|*.pdf", style=wx.FD_SAVE | wx.FD_OVERWRITE_PROMPT) as save_dlg:
                    if save_dlg.ShowModal() == wx.ID_OK:
                        save_path = save_dlg.GetPath()
                        self.pdf_doc.save(save_path)
                        show_msg(_("Annotation added and PDF saved."), _("Success"))
            except Exception as e:
                show_msg(_("Annotation failed: {}").format(str(e)), _("Error"), True)
        else:
            dlg.Destroy()

    def on_rotate_cw(self, event):
        if self.file_type != 'pdf':
            show_msg(_("Rotation is only available for PDF files."), _("Information"))
            return
        self.pdf_doc[self.current_section].set_rotation(90)
        self.pdf_doc.saveIncr()
        self.page_cache.clear()
        self.load_section(self.current_section)
        show_msg(_("Section rotated clockwise."), _("Success"))

    def on_rotate_ccw(self, event):
        if self.file_type != 'pdf':
            show_msg(_("Rotation is only available for PDF files."), _("Information"))
            return
        self.pdf_doc[self.current_section].set_rotation(-90)
        self.pdf_doc.saveIncr()
        self.page_cache.clear()
        self.load_section(self.current_section)
        show_msg(_("Section rotated counter-clockwise."), _("Success"))

    def on_delete_page(self, event):
        if self.file_type != 'pdf':
            show_msg(_("Deletion is only available for PDF files."), _("Information"))
            return
        if self.total_sections <= 1:
            show_msg(_("Cannot delete the only section."), _("Error"), True)
            return
        confirm = wx.MessageDialog(self, _("Delete current section?"), _("Confirm"), wx.YES_NO | wx.ICON_QUESTION)
        if confirm.ShowModal() == wx.ID_YES:
            self.pdf_doc.delete_page(self.current_section)
            self.pdf_doc.saveIncr()
            self.total_sections -= 1
            self.sections.pop(self.current_section)
            self.section_choice.Clear()
            self.section_choice.AppendItems([str(i+1) for i in range(self.total_sections)])
            if self.current_section >= self.total_sections:
                self.current_section = self.total_sections - 1
            self.page_cache.clear()
            self.load_section(self.current_section)
            show_msg(_("Section deleted."), _("Success"))
        confirm.Destroy()

    def on_extract_pages(self, event):
        if self.file_type != 'pdf':
            show_msg(_("Extraction is only available for PDF files."), _("Information"))
            return
        dlg = wx.TextEntryDialog(self, _("Section range (e.g., 1-5,7):"), _("Extract Sections"))
        if dlg.ShowModal() == wx.ID_OK:
            pages = self.parse_section_range(dlg.GetValue())
            dlg.Destroy()
            if not pages:
                show_msg(_("No valid sections."), _("Error"), True)
                return
            default_dir = self.settings.get("export_dir", "")
            with wx.FileDialog(self, _("Save extracted PDF"), defaultDir=default_dir, wildcard="PDF files (*.pdf)|*.pdf", style=wx.FD_SAVE | wx.FD_OVERWRITE_PROMPT) as save_dlg:
                if save_dlg.ShowModal() == wx.ID_OK:
                    save_path = save_dlg.GetPath()
                    new_doc = fitz.open()
                    for p in pages:
                        new_doc.insert_pdf(self.pdf_doc, from_page=p, to_page=p)
                    new_doc.save(save_path)
                    new_doc.close()
                    show_msg(_("Sections extracted successfully."), _("Success"))
        else:
            dlg.Destroy()

    def on_merge_pdf(self, event):
        if self.file_type != 'pdf':
            show_msg(_("Merge is only available for PDF files."), _("Information"))
            return
        with wx.FileDialog(self, _("Select PDF to merge"), wildcard="PDF files (*.pdf)|*.pdf") as file_dlg:
            if file_dlg.ShowModal() == wx.ID_OK:
                other_path = file_dlg.GetPath()
                other_doc = fitz.open(other_path)
                self.pdf_doc.insert_pdf(other_doc)
                other_doc.close()
                self.pdf_doc.saveIncr()
                self.total_sections = len(self.pdf_doc)
                self.sections = [{'type': 'pdf_page', 'page': i} for i in range(self.total_sections)]
                self.section_choice.Clear()
                self.section_choice.AppendItems([str(i+1) for i in range(self.total_sections)])
                self.page_cache.clear()
                self.load_section(self.current_section)
                show_msg(_("PDFs merged successfully."), _("Success"))

    def on_split_pdf(self, event):
        if self.file_type != 'pdf':
            show_msg(_("Split is only available for PDF files."), _("Information"))
            return
        default_dir = self.settings.get("export_dir", "")
        with wx.FileDialog(self, _("Save split PDFs prefix"), defaultDir=default_dir, style=wx.FD_SAVE | wx.FD_OVERWRITE_PROMPT, wildcard="PDF files (*.pdf)|*.pdf") as dlg:
            if dlg.ShowModal() == wx.ID_OK:
                base_path = dlg.GetPath()
                for i in range(self.total_sections):
                    new_doc = fitz.open()
                    new_doc.insert_pdf(self.pdf_doc, from_page=i, to_page=i)
                    out_path = base_path.replace(".pdf", f"_section{i+1}.pdf")
                    new_doc.save(out_path)
                    new_doc.close()
                show_msg(_("PDF split into {} sections.").format(self.total_sections), _("Success"))

    def on_extract_images(self, event):
        if self.file_type != 'pdf':
            show_msg(_("Image extraction is only available for PDF files."), _("Information"))
            return
        default_dir = self.settings.get("export_dir", "")
        with wx.DirDialog(self, _("Select output folder"), defaultPath=default_dir) as dlg:
            if dlg.ShowModal() == wx.ID_OK:
                folder = dlg.GetPath()
                count = 0
                for i in range(self.total_sections):
                    for img in self.pdf_doc[i].get_images():
                        xref = img[0]
                        base_image = self.pdf_doc.extract_image(xref)
                        img_bytes = base_image["image"]
                        ext = base_image["ext"]
                        with open(os.path.join(folder, f"image_section{i+1}_{xref}.{ext}"), "wb") as f:
                            f.write(img_bytes)
                        count += 1
                show_msg(_("Extracted {} images.").format(count), _("Success"))

    def on_save_page_image(self, event):
        if self.file_type != 'pdf':
            show_msg(_("This feature is only available for PDF files."), _("Information"))
            return
        default_dir = self.settings.get("export_dir", "")
        with wx.FileDialog(self, _("Save Section Image"), defaultDir=default_dir, wildcard="PNG files (*.png)|*.png", style=wx.FD_SAVE | wx.FD_OVERWRITE_PROMPT) as dlg:
            if dlg.ShowModal() == wx.ID_OK:
                pix = self.pdf_doc[self.current_section].get_pixmap(dpi=300)
                pix.save(dlg.GetPath())
                show_msg(_("Section saved as image."), _("Success"))

    def on_add_watermark(self, event):
        if self.file_type != 'pdf':
            show_msg(_("Watermark is only available for PDF files."), _("Information"))
            return
        dlg = wx.TextEntryDialog(self, _("Watermark text:"), _("Add Watermark"))
        if dlg.ShowModal() == wx.ID_OK:
            text = dlg.GetValue()
            for i in range(self.total_sections):
                page = self.pdf_doc[i]
                rect = page.rect
                box = fitz.Rect(rect.width*0.1, rect.height*0.4, rect.width*0.9, rect.height*0.6)
                page.insert_textbox(box, text, fontsize=50, color=(0.8, 0.8, 0.8), rotate=45, align=1)
            self.pdf_doc.saveIncr()
            self.page_cache.clear()
            self.load_section(self.current_section)
            show_msg(_("Watermark added."), _("Success"))
        dlg.Destroy()

    def on_redact_text(self, event):
        if self.file_type != 'pdf':
            show_msg(_("Redaction is only available for PDF files."), _("Information"))
            return
        dlg = wx.TextEntryDialog(self, _("Text to redact:"), _("Redact Text"))
        if dlg.ShowModal() == wx.ID_OK:
            text = dlg.GetValue()
            count = 0
            for i in range(self.total_sections):
                page = self.pdf_doc[i]
                rects = page.search_for(text)
                if rects:
                    for rect in rects:
                        page.add_redact_annot(rect)
                    page.apply_redactions()
                    count += len(rects)
            if count > 0:
                self.pdf_doc.saveIncr()
                self.page_cache.clear()
                self.load_section(self.current_section)
                show_msg(_("Redacted {} occurrences.").format(count), _("Success"))
            else:
                show_msg(_("Text not found."), _("Information"))
        dlg.Destroy()

    def save_history(self, section):
        data = load_data()
        data["history"][self.file_path] = section
        data["reading_sessions"].append({
            "file_path": self.file_path,
            "section": section,
            "timestamp": time.time()
        })
        save_data(data)

    def save_recent_file(self):
        data = load_data()
        data["recent_files"][self.file_path] = time.time()
        save_data(data)

    def on_close_dialog(self, event):
        self.save_settings()
        if self.auto_scroll:
            self.stop_auto_scroll()
        if self.read_aloud:
            self.stop_read_aloud()
        if self.pdf_doc:
            self.pdf_doc.close()
        if self.GetParent():
            self.GetParent().Raise()
        self.Destroy()

    def on_text_key_down(self, event):
        event.Skip()

    def show_menu_navigator(self, event=None):
        commands = [
            ("File", self.on_file_menu),
            ("Edit", self.on_edit_menu),
            ("View", self.on_view_menu),
            ("Tools", self.on_tools_menu),
            ("Navigate", self.on_nav_menu),
            ("Manipulations", self.on_manipulations_menu),
            ("Help", self.on_help_menu),
            ("Settings", self.on_settings_dialog),
        ]
        choices = [c[0] for c in commands]
        dlg = wx.SingleChoiceDialog(self, _("Select a menu:"), _("Menu Navigator"), choices)
        if dlg.ShowModal() == wx.ID_OK:
            selection = dlg.GetSelection()
            if 0 <= selection < len(commands):
                commands[selection][1](None)
        dlg.Destroy()

class AnnotationEntryDialog(wx.Dialog):
    def __init__(self, parent):
        super().__init__(parent, title=_("Add Annotation"))
        self.annotation_text = ""
        self.SetSize((400, 250))
        self.Centre()
        main_sizer = wx.BoxSizer(wx.VERTICAL)
        label = wx.StaticText(self, label=_("Annotation text:"))
        main_sizer.Add(label, 0, wx.ALL, 5)
        self.text_ctrl = wx.TextCtrl(self, style=wx.TE_MULTILINE, size=(-1, 100))
        main_sizer.Add(self.text_ctrl, 1, wx.ALL | wx.EXPAND, 5)
        btn_sizer = wx.BoxSizer(wx.HORIZONTAL)
        add_btn = wx.Button(self, label=_("&Add"))
        add_btn.Bind(wx.EVT_BUTTON, self.on_add)
        btn_sizer.Add(add_btn, 0, wx.RIGHT, 10)
        cancel_btn = wx.Button(self, label=_("&Cancel"))
        cancel_btn.Bind(wx.EVT_BUTTON, self.on_cancel)
        btn_sizer.Add(cancel_btn, 0)
        main_sizer.Add(btn_sizer, 0, wx.ALL | wx.CENTER, 10)
        self.SetSizer(main_sizer)
        self.text_ctrl.SetFocus()

    def on_add(self, event):
        self.annotation_text = self.text_ctrl.GetValue()
        if not self.annotation_text.strip():
            show_msg(_("Annotation text cannot be empty."), _("Error"), True)
            return
        self.EndModal(wx.ID_OK)

    def on_cancel(self, event):
        self.EndModal(wx.ID_CANCEL)

class SearchDialog(wx.Dialog):
    def __init__(self, parent):
        super().__init__(parent, title=_("Search Document"))
        self.SetSize((450, 350))
        self.Centre()
        self.query = ""
        self.case_sensitive = False
        self.whole_words = False
        self.regex = False
        self.search_range = "all"

        main_sizer = wx.BoxSizer(wx.VERTICAL)
        input_sizer = wx.BoxSizer(wx.HORIZONTAL)
        lbl = wx.StaticText(self, label=_("Find:"))
        input_sizer.Add(lbl, 0, wx.RIGHT | wx.CENTER, 5)
        self.text_ctrl = wx.TextCtrl(self)
        input_sizer.Add(self.text_ctrl, 1, wx.EXPAND)
        main_sizer.Add(input_sizer, 0, wx.ALL | wx.EXPAND, 10)

        self.case_cb = wx.CheckBox(self, label=_("&Case sensitive"))
        main_sizer.Add(self.case_cb, 0, wx.ALL, 5)
        self.whole_cb = wx.CheckBox(self, label=_("&Whole words"))
        main_sizer.Add(self.whole_cb, 0, wx.ALL, 5)
        self.regex_cb = wx.CheckBox(self, label=_("Regular e&xpression"))
        main_sizer.Add(self.regex_cb, 0, wx.ALL, 5)

        range_label = wx.StaticText(self, label=_("Search scope:"))
        main_sizer.Add(range_label, 0, wx.ALL, 5)
        self.range_choice = wx.Choice(self, choices=[_("All sections"), _("Current section")])
        self.range_choice.SetSelection(0)
        main_sizer.Add(self.range_choice, 0, wx.ALL | wx.EXPAND, 5)

        data = load_data()
        history = data.get("search_history", [])
        if history:
            history_label = wx.StaticText(self, label=_("Recent searches:"))
            main_sizer.Add(history_label, 0, wx.ALL, 5)
            self.history_choice = wx.Choice(self, choices=history[:10])
            self.history_choice.Bind(wx.EVT_CHOICE, self.on_history_select)
            main_sizer.Add(self.history_choice, 0, wx.ALL | wx.EXPAND, 5)

        btn_sizer = wx.BoxSizer(wx.HORIZONTAL)
        ok_btn = wx.Button(self, label=_("&Search"))
        ok_btn.Bind(wx.EVT_BUTTON, self.on_ok)
        btn_sizer.Add(ok_btn, 0, wx.RIGHT, 10)
        cancel_btn = wx.Button(self, label=_("&Cancel"))
        cancel_btn.Bind(wx.EVT_BUTTON, self.on_cancel)
        btn_sizer.Add(cancel_btn, 0)
        main_sizer.Add(btn_sizer, 0, wx.ALL | wx.CENTER, 10)
        self.SetSizer(main_sizer)
        self.text_ctrl.SetFocus()

    def on_history_select(self, event):
        self.text_ctrl.SetValue(self.history_choice.GetStringSelection())

    def on_ok(self, event):
        self.query = self.text_ctrl.GetValue()
        self.case_sensitive = self.case_cb.IsChecked()
        self.whole_words = self.whole_cb.IsChecked()
        self.regex = self.regex_cb.IsChecked()
        self.search_range = "current" if self.range_choice.GetSelection() == 1 else "all"
        self.EndModal(wx.ID_OK)

    def on_cancel(self, event):
        self.EndModal(wx.ID_CANCEL)

class SearchResultsDialog(wx.Dialog):
    def __init__(self, parent, results):
        super().__init__(parent, title=_("Search Results"))
        self.results = results
        self.result_index = -1
        self.SetSize((600, 450))
        self.Centre()
        main_sizer = wx.BoxSizer(wx.VERTICAL)
        self.list_box = wx.ListBox(self, style=wx.LB_SINGLE)
        for r in results:
            section, snippet, _, match_text = r
            self.list_box.Append(_("Section {}: {}").format(section+1, match_text))
        self.list_box.Bind(wx.EVT_LISTBOX_DCLICK, self.on_select)
        self.list_box.Bind(wx.EVT_LISTBOX, self.on_list_select)
        main_sizer.Add(self.list_box, 1, wx.ALL | wx.EXPAND, 10)

        self.preview = wx.TextCtrl(self, style=wx.TE_MULTILINE | wx.TE_READONLY, size=(-1, 60))
        main_sizer.Add(self.preview, 0, wx.ALL | wx.EXPAND, 10)

        btn_sizer = wx.BoxSizer(wx.HORIZONTAL)
        ok_btn = wx.Button(self, label=_("&Jump"))
        ok_btn.Bind(wx.EVT_BUTTON, self.on_ok)
        btn_sizer.Add(ok_btn, 0, wx.RIGHT, 10)
        cancel_btn = wx.Button(self, label=_("&Cancel"))
        cancel_btn.Bind(wx.EVT_BUTTON, self.on_cancel)
        btn_sizer.Add(cancel_btn, 0)
        main_sizer.Add(btn_sizer, 0, wx.ALL | wx.CENTER, 10)
        self.SetSizer(main_sizer)
        if self.list_box.GetCount() > 0:
            self.list_box.SetSelection(0)
            self.list_box.SetFocus()
            self.on_list_select(None)

    def on_list_select(self, event):
        idx = self.list_box.GetSelection()
        if idx >= 0:
            _, snippet, _, _ = self.results[idx]
            self.preview.SetValue(snippet)

    def on_select(self, event):
        self.on_ok(event)

    def on_ok(self, event):
        self.result_index = self.list_box.GetSelection()
        self.EndModal(wx.ID_OK)

    def on_cancel(self, event):
        self.EndModal(wx.ID_CANCEL)

class BookmarksManagerDialog(wx.Dialog):
    def __init__(self, parent, bookmarks):
        super().__init__(parent, title=_("Bookmarks Manager"))
        self.bookmarks = bookmarks
        self.action = ""
        self.selected_section = None
        self.selected_id = None
        self.new_title = ""
        self.SetSize((550, 450))
        self.Centre()
        main_sizer = wx.BoxSizer(wx.VERTICAL)
        self.list_box = wx.ListBox(self, style=wx.LB_SINGLE)
        for b in bookmarks:
            self.list_box.Append(_("Section {}: {}").format(b["section"]+1, b["title"]))
        main_sizer.Add(self.list_box, 1, wx.ALL | wx.EXPAND, 10)

        self.detail_text = wx.TextCtrl(self, style=wx.TE_MULTILINE | wx.TE_READONLY, size=(-1, 40))
        main_sizer.Add(self.detail_text, 0, wx.ALL | wx.EXPAND, 5)

        btn_sizer = wx.BoxSizer(wx.HORIZONTAL)
        jump_btn = wx.Button(self, label=_("&Jump to"))
        jump_btn.Bind(wx.EVT_BUTTON, self.on_jump)
        btn_sizer.Add(jump_btn, 0, wx.RIGHT, 5)
        rename_btn = wx.Button(self, label=_("Rena&me"))
        rename_btn.Bind(wx.EVT_BUTTON, self.on_rename)
        btn_sizer.Add(rename_btn, 0, wx.RIGHT, 5)
        delete_btn = wx.Button(self, label=_("&Delete"))
        delete_btn.Bind(wx.EVT_BUTTON, self.on_delete)
        btn_sizer.Add(delete_btn, 0, wx.RIGHT, 5)
        export_btn = wx.Button(self, label=_("E&xport"))
        export_btn.Bind(wx.EVT_BUTTON, self.on_export)
        btn_sizer.Add(export_btn, 0, wx.RIGHT, 5)
        close_btn = wx.Button(self, label=_("&Close"))
        close_btn.Bind(wx.EVT_BUTTON, self.on_close)
        btn_sizer.Add(close_btn, 0)
        main_sizer.Add(btn_sizer, 0, wx.ALL | wx.CENTER, 10)
        self.SetSizer(main_sizer)
        if self.list_box.GetCount() > 0:
            self.list_box.SetSelection(0)
            self.update_detail()
        self.list_box.SetFocus()
        self.list_box.Bind(wx.EVT_LISTBOX, self.on_list_select)

    def on_list_select(self, event):
        self.update_detail()

    def update_detail(self):
        idx = self.list_box.GetSelection()
        if idx >= 0:
            b = self.bookmarks[idx]
            self.detail_text.SetValue(_("Section: {}\nTitle: {}\nAdded: {}").format(
                b["section"]+1, b["title"], time.strftime("%Y-%m-%d %H:%M", time.localtime(b["timestamp"]))
            ))

    def get_selected(self):
        idx = self.list_box.GetSelection()
        if idx >= 0:
            self.selected_id = self.bookmarks[idx]["id"]
            self.selected_section = self.bookmarks[idx]["section"]
            return True
        show_msg(_("No bookmark selected"), _("Information"))
        return False

    def on_jump(self, event):
        if self.get_selected():
            self.action = "jump"
            self.EndModal(wx.ID_OK)

    def on_rename(self, event):
        if self.get_selected():
            dlg = wx.TextEntryDialog(self, _("New title:"), _("Rename Bookmark"))
            if dlg.ShowModal() == wx.ID_OK:
                self.new_title = dlg.GetValue()
                self.action = "rename"
                dlg.Destroy()
                self.EndModal(wx.ID_OK)
            else:
                dlg.Destroy()

    def on_delete(self, event):
        if self.get_selected():
            self.action = "delete"
            self.EndModal(wx.ID_OK)

    def on_export(self, event):
        with wx.FileDialog(self, _("Export bookmarks"), wildcard="JSON files (*.json)|*.json", style=wx.FD_SAVE | wx.FD_OVERWRITE_PROMPT) as dlg:
            if dlg.ShowModal() == wx.ID_OK:
                path = dlg.GetPath()
                with open(path, "w", encoding="utf-8") as f:
                    json.dump(self.bookmarks, f, ensure_ascii=False, indent=2)
                show_msg(_("Bookmarks exported."), _("Success"))

    def on_close(self, event):
        self.EndModal(wx.ID_CANCEL)

class TocDialog(wx.Dialog):
    def __init__(self, parent, toc, file_type):
        super().__init__(parent, title=_("Table of Contents"))
        self.file_type = file_type
        self.toc = toc
        self.selected_page = None
        self.selected_section = None
        self.SetSize((600, 500))
        self.Centre()
        main_sizer = wx.BoxSizer(wx.VERTICAL)
        if file_type == 'pdf':
            self.tree = wx.TreeCtrl(self, style=wx.TR_DEFAULT_STYLE | wx.TR_HIDE_ROOT)
            root = self.tree.AddRoot("TOC")
            self.add_pdf_items(root, toc)
            self.tree.ExpandAll()
            self.tree.Bind(wx.EVT_TREE_ITEM_ACTIVATED, self.on_activate_pdf)
            main_sizer.Add(self.tree, 1, wx.ALL | wx.EXPAND, 10)
        elif file_type == 'docx':
            self.list_box = wx.ListBox(self)
            for item in toc:
                self.list_box.Append(_("{} (Section {})").format(item[1], item[2]))
            self.list_box.Bind(wx.EVT_LISTBOX_DCLICK, self.on_activate_docx)
            main_sizer.Add(self.list_box, 1, wx.ALL | wx.EXPAND, 10)
        elif file_type == 'epub':
            self.list_box = wx.ListBox(self)
            for item in toc:
                self.list_box.Append(_("{} (Section {})").format(item[1], item[2]))
            self.list_box.Bind(wx.EVT_LISTBOX_DCLICK, self.on_activate_epub)
            main_sizer.Add(self.list_box, 1, wx.ALL | wx.EXPAND, 10)

        btn_sizer = wx.BoxSizer(wx.HORIZONTAL)
        jump_btn = wx.Button(self, label=_("&Jump"))
        if file_type == 'pdf':
            jump_btn.Bind(wx.EVT_BUTTON, self.on_jump_pdf)
        elif file_type == 'docx':
            jump_btn.Bind(wx.EVT_BUTTON, self.on_jump_docx)
        elif file_type == 'epub':
            jump_btn.Bind(wx.EVT_BUTTON, self.on_jump_epub)
        btn_sizer.Add(jump_btn, 0, wx.RIGHT, 10)
        close_btn = wx.Button(self, label=_("&Close"))
        close_btn.Bind(wx.EVT_BUTTON, self.on_close)
        btn_sizer.Add(close_btn, 0)
        main_sizer.Add(btn_sizer, 0, wx.ALL | wx.CENTER, 10)
        self.SetSizer(main_sizer)
        if file_type == 'pdf' and self.tree.GetCount():
            self.tree.SetFocus()

    def add_pdf_items(self, parent, items):
        for item in items:
            level, title, page = item[0], item[1], item[2]
            indent = "  " * (level - 1)
            child = self.tree.AppendItem(parent, _("{}{} (Page {})").format(indent, title, page))
            self.tree.SetItemData(child, {"page": page, "title": title, "level": level})

    def on_activate_pdf(self, event):
        item = self.tree.GetSelection()
        data = self.tree.GetItemData(item)
        if data:
            self.selected_page = data["page"]
            self.EndModal(wx.ID_OK)

    def on_jump_pdf(self, event):
        item = self.tree.GetSelection()
        if item:
            data = self.tree.GetItemData(item)
            if data:
                self.selected_page = data["page"]
                self.EndModal(wx.ID_OK)
            else:
                show_msg(_("No selection"), _("Information"))
        else:
            show_msg(_("No item selected"), _("Information"))

    def on_activate_docx(self, event):
        idx = self.list_box.GetSelection()
        if idx >= 0:
            self.selected_section = self.toc[idx][2]
            self.EndModal(wx.ID_OK)

    def on_jump_docx(self, event):
        idx = self.list_box.GetSelection()
        if idx >= 0:
            self.selected_section = self.toc[idx][2]
            self.EndModal(wx.ID_OK)
        else:
            show_msg(_("No item selected"), _("Information"))

    def on_activate_epub(self, event):
        idx = self.list_box.GetSelection()
        if idx >= 0:
            self.selected_section = self.toc[idx][2]
            self.EndModal(wx.ID_OK)

    def on_jump_epub(self, event):
        idx = self.list_box.GetSelection()
        if idx >= 0:
            self.selected_section = self.toc[idx][2]
            self.EndModal(wx.ID_OK)
        else:
            show_msg(_("No item selected"), _("Information"))

    def on_close(self, event):
        self.EndModal(wx.ID_CANCEL)

class OCRDialog(wx.Dialog):
    def __init__(self, parent):
        super().__init__(parent, title=_("OCR Options"))
        self.SetSize((350, 250))
        self.Centre()
        main_sizer = wx.BoxSizer(wx.VERTICAL)
        self.cur_radio = wx.RadioButton(self, label=_("&Current section"), style=wx.RB_GROUP)
        self.all_radio = wx.RadioButton(self, label=_("&All sections"))
        self.range_radio = wx.RadioButton(self, label=_("&Section range"))
        main_sizer.Add(self.cur_radio, 0, wx.ALL, 10)
        main_sizer.Add(self.all_radio, 0, wx.ALL, 10)
        main_sizer.Add(self.range_radio, 0, wx.ALL, 10)

        range_sizer = wx.BoxSizer(wx.HORIZONTAL)
        range_sizer.Add(wx.StaticText(self, label=_("From:")), 0, wx.RIGHT | wx.CENTER, 5)
        self.from_text = wx.TextCtrl(self, size=(60, -1))
        range_sizer.Add(self.from_text, 0, wx.RIGHT, 10)
        range_sizer.Add(wx.StaticText(self, label=_("To:")), 0, wx.RIGHT | wx.CENTER, 5)
        self.to_text = wx.TextCtrl(self, size=(60, -1))
        range_sizer.Add(self.to_text, 0)
        main_sizer.Add(range_sizer, 0, wx.ALL, 10)

        self.cur_radio.SetValue(True)

        btn_sizer = wx.BoxSizer(wx.HORIZONTAL)
        ok_btn = wx.Button(self, label=_("&Start"))
        ok_btn.Bind(wx.EVT_BUTTON, self.on_ok)
        btn_sizer.Add(ok_btn, 0, wx.RIGHT, 10)
        cancel_btn = wx.Button(self, label=_("&Cancel"))
        cancel_btn.Bind(wx.EVT_BUTTON, self.on_cancel)
        btn_sizer.Add(cancel_btn, 0)
        main_sizer.Add(btn_sizer, 0, wx.ALL | wx.CENTER, 10)
        self.SetSizer(main_sizer)
        self.cur_radio.SetFocus()

    def get_pages(self):
        parent = self.GetParent()
        if self.cur_radio.GetValue():
            return [parent.current_section]
        elif self.all_radio.GetValue():
            return list(range(parent.total_sections))
        else:
            try:
                start = int(self.from_text.GetValue()) - 1
                end = int(self.to_text.GetValue())
                return list(range(max(0, start), min(parent.total_sections, end)))
            except:
                return [parent.current_section]

    def on_ok(self, event):
        self.EndModal(wx.ID_OK)

    def on_cancel(self, event):
        self.EndModal(wx.ID_CANCEL)

class ExportDialog(wx.Dialog):
    def __init__(self, parent):
        super().__init__(parent, title=_("Export Options"))
        self.SetSize((380, 350))
        self.Centre()
        self.export_type = "txt"
        self.scope = "current"
        self.include_notes = False
        self.include_ocr = False
        main_sizer = wx.BoxSizer(wx.VERTICAL)

        type_label = wx.StaticText(self, label=_("Format:"))
        main_sizer.Add(type_label, 0, wx.ALL, 5)
        self.type_choice = wx.Choice(self, choices=[_("Text (TXT)"), _("Word (DOCX)"), _("Audiobook (MP3)"), _("CSV (Tables)"), _("JSON"), _("HTML")])
        self.type_choice.SetSelection(0)
        main_sizer.Add(self.type_choice, 0, wx.ALL | wx.EXPAND, 5)

        scope_label = wx.StaticText(self, label=_("Scope:"))
        main_sizer.Add(scope_label, 0, wx.ALL, 5)
        self.cur_radio = wx.RadioButton(self, label=_("&Current section"), style=wx.RB_GROUP)
        self.all_radio = wx.RadioButton(self, label=_("&Entire document"))
        self.tagged_radio = wx.RadioButton(self, label=_("&Tagged sections"))
        main_sizer.Add(self.cur_radio, 0, wx.ALL, 5)
        main_sizer.Add(self.all_radio, 0, wx.ALL, 5)
        main_sizer.Add(self.tagged_radio, 0, wx.ALL, 5)

        self.notes_cb = wx.CheckBox(self, label=_("Include ¬es"))
        main_sizer.Add(self.notes_cb, 0, wx.ALL, 5)

        self.ocr_cb = wx.CheckBox(self, label=_("Include O&CR text"))
        main_sizer.Add(self.ocr_cb, 0, wx.ALL, 5)

        btn_sizer = wx.BoxSizer(wx.HORIZONTAL)
        ok_btn = wx.Button(self, label=_("&Export"))
        ok_btn.Bind(wx.EVT_BUTTON, self.on_ok)
        btn_sizer.Add(ok_btn, 0, wx.RIGHT, 10)
        cancel_btn = wx.Button(self, label=_("&Cancel"))
        cancel_btn.Bind(wx.EVT_BUTTON, self.on_cancel)
        btn_sizer.Add(cancel_btn, 0)
        main_sizer.Add(btn_sizer, 0, wx.ALL | wx.CENTER, 10)
        self.SetSizer(main_sizer)
        self.type_choice.SetFocus()

    def on_ok(self, event):
        type_idx = self.type_choice.GetSelection()
        self.export_type = ["txt", "docx", "audiobook", "csv", "json", "html"][type_idx]
        if self.cur_radio.GetValue():
            self.scope = "current"
        elif self.tagged_radio.GetValue():
            self.scope = "tagged"
        else:
            self.scope = "all"
        self.include_notes = self.notes_cb.IsChecked()
        self.include_ocr = self.ocr_cb.IsChecked()
        self.EndModal(wx.ID_OK)

    def on_cancel(self, event):
        self.EndModal(wx.ID_CANCEL)

class MetadataDialog(wx.Dialog):
    def __init__(self, parent, file_path, pdf_doc):
        super().__init__(parent, title=_("Document Information"))
        self.SetSize((550, 450))
        self.Centre()
        meta = pdf_doc.metadata
        info = []
        info.append(_("Title: {}").format(meta.get("title", _("Not available"))))
        info.append(_("Author: {}").format(meta.get("author", _("Not available"))))
        info.append(_("Subject: {}").format(meta.get("subject", _("Not available"))))
        info.append(_("Keywords: {}").format(meta.get("keywords", _("Not available"))))
        info.append(_("Creator: {}").format(meta.get("creator", _("Not available"))))
        info.append(_("Producer: {}").format(meta.get("producer", _("Not available"))))
        info.append(_("Creation Date: {}").format(meta.get("creationDate", _("Not available"))))
        info.append(_("Modification Date: {}").format(meta.get("modDate", _("Not available"))))
        info.append(_("Format: {}").format(meta.get("format", "PDF")))
        info.append(_("Encryption: {}").format(meta.get("encryption", _("None"))))
        info.append(_("Page Count: {}").format(len(pdf_doc)))
        info.append(_("File Size: {:.2f} MB").format(os.path.getsize(file_path) / (1024*1024)))
        main_sizer = wx.BoxSizer(wx.VERTICAL)
        text_ctrl = wx.TextCtrl(self, style=wx.TE_MULTILINE | wx.TE_READONLY, value="\n".join(info))
        main_sizer.Add(text_ctrl, 1, wx.ALL | wx.EXPAND, 10)
        btn_sizer = wx.BoxSizer(wx.HORIZONTAL)
        copy_btn = wx.Button(self, label=_("&Copy Information"))
        copy_btn.Bind(wx.EVT_BUTTON, lambda e: self.copy_info(text_ctrl.GetValue()))
        btn_sizer.Add(copy_btn, 0, wx.RIGHT, 10)
        close_btn = wx.Button(self, label=_("&Close"))
        close_btn.Bind(wx.EVT_BUTTON, self.on_close)
        btn_sizer.Add(close_btn, 0)
        main_sizer.Add(btn_sizer, 0, wx.ALL | wx.CENTER, 10)
        self.SetSizer(main_sizer)
        text_ctrl.SetFocus()

    def copy_info(self, text):
        if wx.TheClipboard.Open():
            wx.TheClipboard.SetData(wx.TextDataObject(text))
            wx.TheClipboard.Close()
            show_msg(_("Metadata copied to clipboard"), _("Success"))

    def on_close(self, event):
        self.EndModal(wx.ID_OK)

class AboutDialog(wx.Dialog):
    def __init__(self, parent):
        super().__init__(parent, title=_("About Document Reader"))
        self.SetSize((450, 250))
        self.Centre()
        main_sizer = wx.BoxSizer(wx.VERTICAL)
        msg = _("Document Reader for NVDA\nVersion 3.0\nDesigned for blind and visually impaired users.\nFeatures: OCR, TTS, Search, Bookmarks, Notes, Export\nTelegram: @blindtechvisionary")
        label = wx.StaticText(self, label=msg)
        label.Wrap(400)
        main_sizer.Add(label, 0, wx.ALL | wx.CENTER, 15)
        telegram_btn = wx.Button(self, label=_("Join &Telegram"))
        telegram_btn.Bind(wx.EVT_BUTTON, lambda e: webbrowser.open("https://t.me/blindtechvisionary"))
        main_sizer.Add(telegram_btn, 0, wx.ALL | wx.CENTER, 10)
        back_btn = wx.Button(self, label=_("&Back"))
        back_btn.Bind(wx.EVT_BUTTON, self.on_back)
        main_sizer.Add(back_btn, 0, wx.ALL | wx.CENTER, 10)
        self.SetSizer(main_sizer)
        telegram_btn.SetFocus()

    def on_back(self, event):
        self.EndModal(wx.ID_OK)
        self.Destroy()

class HelpDialog(wx.Dialog):
    def __init__(self, parent):
        super().__init__(parent, title=_("Help"))
        self.SetSize((600, 500))
        self.Centre()
        main_sizer = wx.BoxSizer(wx.VERTICAL)
        help_text = _(
            "Keyboard Shortcuts:\n\n"
            "Navigation:\n"
            "Alt+N / Alt+PageDown: Next Section\n"
            "Alt+P / Alt+PageUp: Previous Section\n"
            "Ctrl+G: Go To Section\n"
            "Ctrl+T: Table of Contents\n"
            "Ctrl+B: Add Bookmark\n"
            "Ctrl+Shift+B: Bookmarks Manager\n\n"
            "Search:\n"
            "Ctrl+F: Search\n"
            "F3: Next Search Result\n"
            "Shift+F3: Previous Search Result\n\n"
            "View:\n"
            "Ctrl++: Zoom In\n"
            "Ctrl+-: Zoom Out\n"
            "Ctrl+0: Reset Zoom\n"
            "Ctrl+H: High Contrast\n"
            "Ctrl+Shift+I: Invert Colors\n"
            "Ctrl+Shift+V: Split View\n\n"
            "Tools:\n"
            "Ctrl+R: Read Aloud\n"
            "Ctrl+Shift+A: Auto Scroll\n"
            "Ctrl+Shift+O: OCR\n"
            "Ctrl+Shift+P: Extraction Profile\n"
            "Ctrl+Shift+S: Statistics\n"
            "Ctrl+Shift+C: Compare Sections\n"
            "Ctrl+Shift+L: Highlight Selection\n"
            "Ctrl+Shift+T: Translate Section\n\n"
            "Manipulations:\n"
            "Ctrl+Shift+M: Menu Navigator\n"
            "Add Annotation, Rotate, Delete, Extract, Merge, Split\n\n"
            "Editing:\n"
            "Ctrl+C: Copy Section\n"
            "Ctrl+Shift+C: Copy All\n\n"
            "Export:\n"
            "Ctrl+Shift+E: Export Menu\n\n"
            "Other:\n"
            "Ctrl+M: Properties\n"
            "Ctrl+I: Import Document\n"
            "Ctrl+W: Close Viewer\n"
            "F1: This Help\n"
            "Ctrl+F1: About\n"
            "Escape: Close Current Dialog\n\n"
            "Reading modes change text extraction method.\n"
            "Use extraction profiles for custom text processing."
        )
        text_ctrl = wx.TextCtrl(self, style=wx.TE_MULTILINE | wx.TE_READONLY, value=help_text)
        main_sizer.Add(text_ctrl, 1, wx.ALL | wx.EXPAND, 10)
        close_btn = wx.Button(self, label=_("&Close"))
        close_btn.Bind(wx.EVT_BUTTON, self.on_close)
        main_sizer.Add(close_btn, 0, wx.ALL | wx.CENTER, 10)
        self.SetSizer(main_sizer)
        text_ctrl.SetFocus()

    def on_close(self, event):
        self.EndModal(wx.ID_OK)

class AnnotationsDialog(wx.Dialog):
    def __init__(self, parent, annotations):
        super().__init__(parent, title=_("Annotations"))
        self.SetSize((500, 400))
        self.Centre()
        main_sizer = wx.BoxSizer(wx.VERTICAL)
        self.list_box = wx.ListBox(self, style=wx.LB_SINGLE)
        for a in annotations:
            self.list_box.Append(_("Section {}: {}").format(a["section"]+1, a.get("text", "")[:50]))
        main_sizer.Add(self.list_box, 1, wx.ALL | wx.EXPAND, 10)
        btn_sizer = wx.BoxSizer(wx.HORIZONTAL)
        close_btn = wx.Button(self, label=_("&Close"))
        close_btn.Bind(wx.EVT_BUTTON, self.on_close)
        btn_sizer.Add(close_btn, 0)
        main_sizer.Add(btn_sizer, 0, wx.ALL | wx.CENTER, 10)
        self.SetSizer(main_sizer)

    def on_close(self, event):
        self.EndModal(wx.ID_OK)

class HighlightsDialog(wx.Dialog):
    def __init__(self, parent, highlights, pdf_doc=None, file_path=None):
        super().__init__(parent, title=_("Highlights"))
        self.SetSize((500, 450))
        self.Centre()
        self.highlights = highlights
        self.pdf_doc = pdf_doc
        self.file_path = file_path
        main_sizer = wx.BoxSizer(wx.VERTICAL)
        self.list_box = wx.ListBox(self, style=wx.LB_SINGLE)
        for h in highlights:
            self.list_box.Append(_("Section {}: {}").format(h["section"]+1, h.get("text", "")[:50]))
        main_sizer.Add(self.list_box, 1, wx.ALL | wx.EXPAND, 10)
        btn_sizer = wx.BoxSizer(wx.HORIZONTAL)
        delete_btn = wx.Button(self, label=_("&Delete Highlight"))
        delete_btn.Bind(wx.EVT_BUTTON, self.on_delete)
        btn_sizer.Add(delete_btn, 0, wx.RIGHT, 10)
        close_btn = wx.Button(self, label=_("&Close"))
        close_btn.Bind(wx.EVT_BUTTON, self.on_close)
        btn_sizer.Add(close_btn, 0)
        main_sizer.Add(btn_sizer, 0, wx.ALL | wx.CENTER, 10)
        self.SetSizer(main_sizer)
        if self.list_box.GetCount() > 0:
            self.list_box.SetSelection(0)

    def on_delete(self, event):
        idx = self.list_box.GetSelection()
        if idx < 0:
            show_msg(_("No highlight selected."), _("Information"))
            return
        hl = self.highlights[idx]
        data = load_data()
        data["highlights"] = [h for h in data["highlights"] if h["id"] != hl["id"]]
        save_data(data)
        show_msg(_("Highlight removed from list."), _("Success"))
        self.EndModal(wx.ID_OK)

    def on_close(self, event):
        self.EndModal(wx.ID_CANCEL)

class ExtractionProfileDialog(wx.Dialog):
    def __init__(self, parent):
        super().__init__(parent, title=_("Extraction Profile"))
        self.SetSize((400, 300))
        self.Centre()
        self.selected_profile = "default"
        main_sizer = wx.BoxSizer(wx.VERTICAL)
        profiles = ["default", "simple", "detailed", "ocr_only", "clean"]
        self.profile_choice = wx.Choice(self, choices=profiles)
        self.profile_choice.SetSelection(0)
        main_sizer.Add(self.profile_choice, 0, wx.ALL | wx.EXPAND, 10)

        desc_text = _("Default: Standard text extraction\nSimple: Basic text only\nDetailed: With metadata\nOCR Only: OCR results only\nClean: Remove extra whitespace")
        desc_label = wx.StaticText(self, label=desc_text)
        desc_label.Wrap(350)
        main_sizer.Add(desc_label, 0, wx.ALL, 10)

        btn_sizer = wx.BoxSizer(wx.HORIZONTAL)
        ok_btn = wx.Button(self, label=_("&Apply"))
        ok_btn.Bind(wx.EVT_BUTTON, self.on_ok)
        btn_sizer.Add(ok_btn, 0, wx.RIGHT, 10)
        cancel_btn = wx.Button(self, label=_("&Cancel"))
        cancel_btn.Bind(wx.EVT_BUTTON, self.on_cancel)
        btn_sizer.Add(cancel_btn, 0)
        main_sizer.Add(btn_sizer, 0, wx.ALL | wx.CENTER, 10)
        self.SetSizer(main_sizer)

    def on_ok(self, event):
        self.selected_profile = self.profile_choice.GetStringSelection()
        self.EndModal(wx.ID_OK)

    def on_cancel(self, event):
        self.EndModal(wx.ID_CANCEL)

class SettingsDialog(wx.Dialog):
    def __init__(self, parent):
        super().__init__(parent, title=_("Settings"))
        self.SetSize((600, 500))
        self.Centre()
        self.parent = parent
        self.temp_settings = {
            "voice": parent.settings.get("voice", "en-US-AriaNeural"),
            "tts_speed": parent.settings.get("tts_speed", "+0%"),
            "tts_pitch": parent.settings.get("tts_pitch", "+0Hz"),
            "zoom": parent.zoom_level,
            "contrast": parent.contrast_mode,
            "invert": parent.invert_colors,
            "font_size": parent.font_size,
            "reading_speed": parent.reading_speed,
            "extraction_profile": parent.extraction_profile,
            "export_dir": parent.settings.get("export_dir", ""),
            "tts_volume": parent.settings.get("tts_volume", 100),
            "tts_chunking": parent.settings.get("tts_chunking", True),
        }

        notebook = wx.Notebook(self)
        self.general_panel = self.create_general_panel(notebook)
        notebook.AddPage(self.general_panel, _("General"))
        self.tts_panel = self.create_tts_panel(notebook)
        notebook.AddPage(self.tts_panel, _("Text-to-Speech"))
        self.advanced_panel = self.create_advanced_panel(notebook)
        notebook.AddPage(self.advanced_panel, _("Advanced"))

        main_sizer = wx.BoxSizer(wx.VERTICAL)
        main_sizer.Add(notebook, 1, wx.EXPAND | wx.ALL, 10)

        btn_sizer = wx.BoxSizer(wx.HORIZONTAL)
        ok_btn = wx.Button(self, label=_("&OK"))
        ok_btn.Bind(wx.EVT_BUTTON, self.on_ok)
        btn_sizer.Add(ok_btn, 0, wx.RIGHT, 10)
        cancel_btn = wx.Button(self, label=_("&Cancel"))
        cancel_btn.Bind(wx.EVT_BUTTON, self.on_cancel)
        btn_sizer.Add(cancel_btn, 0)
        main_sizer.Add(btn_sizer, 0, wx.ALL | wx.CENTER, 10)
        self.SetSizer(main_sizer)

        self.update_tts_voice_list()

    def create_general_panel(self, parent):
        panel = wx.Panel(parent)
        sizer = wx.BoxSizer(wx.VERTICAL)
        zoom_sizer = wx.BoxSizer(wx.HORIZONTAL)
        zoom_sizer.Add(wx.StaticText(panel, label=_("Zoom %:")), 0, wx.CENTER)
        self.zoom_spin = wx.SpinCtrl(panel, value=str(self.temp_settings["zoom"]), min=10, max=500)
        self.zoom_spin.SetValue(self.temp_settings["zoom"])
        zoom_sizer.Add(self.zoom_spin, 0)
        sizer.Add(zoom_sizer, 0, wx.ALL, 5)

        font_sizer = wx.BoxSizer(wx.HORIZONTAL)
        font_sizer.Add(wx.StaticText(panel, label=_("Font size:")), 0, wx.CENTER)
        self.font_spin = wx.SpinCtrl(panel, value=str(self.temp_settings["font_size"]), min=6, max=48)
        self.font_spin.SetValue(self.temp_settings["font_size"])
        font_sizer.Add(self.font_spin, 0)
        sizer.Add(font_sizer, 0, wx.ALL, 5)

        self.contrast_cb = wx.CheckBox(panel, label=_("High contrast"))
        self.contrast_cb.SetValue(self.temp_settings["contrast"])
        sizer.Add(self.contrast_cb, 0, wx.ALL, 5)

        self.invert_cb = wx.CheckBox(panel, label=_("Invert colors"))
        self.invert_cb.SetValue(self.temp_settings["invert"])
        sizer.Add(self.invert_cb, 0, wx.ALL, 5)

        speed_sizer = wx.BoxSizer(wx.HORIZONTAL)
        speed_sizer.Add(wx.StaticText(panel, label=_("Auto-scroll speed (0=slow, 9=fast):")), 0, wx.CENTER)
        self.speed_spin = wx.SpinCtrl(panel, value=str(self.temp_settings["reading_speed"]), min=0, max=9)
        self.speed_spin.SetValue(self.temp_settings["reading_speed"])
        speed_sizer.Add(self.speed_spin, 0)
        sizer.Add(speed_sizer, 0, wx.ALL, 5)

        profile_sizer = wx.BoxSizer(wx.HORIZONTAL)
        profile_sizer.Add(wx.StaticText(panel, label=_("Extraction profile:")), 0, wx.CENTER)
        self.profile_choice = wx.Choice(panel, choices=["default", "simple", "detailed", "ocr_only", "clean"])
        self.profile_choice.SetStringSelection(self.temp_settings["extraction_profile"])
        profile_sizer.Add(self.profile_choice, 0)
        sizer.Add(profile_sizer, 0, wx.ALL, 5)

        panel.SetSizer(sizer)
        return panel

    def create_tts_panel(self, parent):
        panel = wx.Panel(parent)
        sizer = wx.BoxSizer(wx.VERTICAL)

        voice_sizer = wx.BoxSizer(wx.HORIZONTAL)
        voice_sizer.Add(wx.StaticText(panel, label=_("Voice:")), 0, wx.CENTER)
        self.voice_choice = wx.Choice(panel, choices=[])
        self.voice_choice.SetStringSelection(self.temp_settings["voice"])
        voice_sizer.Add(self.voice_choice, 1, wx.EXPAND)
        sizer.Add(voice_sizer, 0, wx.ALL | wx.EXPAND, 5)

        fetch_btn = wx.Button(panel, label=_("&Fetch Latest Voices"))
        fetch_btn.Bind(wx.EVT_BUTTON, self.on_fetch_voices)
        sizer.Add(fetch_btn, 0, wx.ALL, 5)

        speed_sizer = wx.BoxSizer(wx.HORIZONTAL)
        speed_sizer.Add(wx.StaticText(panel, label=_("Speed:")), 0, wx.CENTER)
        speeds = ["-50%", "-30%", "-20%", "-10%", "+0%", "+10%", "+20%", "+30%", "+50%"]
        self.speed_choice = wx.Choice(panel, choices=speeds)
        self.speed_choice.SetStringSelection(self.temp_settings["tts_speed"])
        speed_sizer.Add(self.speed_choice, 0)
        sizer.Add(speed_sizer, 0, wx.ALL, 5)

        pitch_sizer = wx.BoxSizer(wx.HORIZONTAL)
        pitch_sizer.Add(wx.StaticText(panel, label=_("Pitch:")), 0, wx.CENTER)
        pitches = ["-20Hz", "-10Hz", "+0Hz", "+10Hz", "+20Hz"]
        self.pitch_choice = wx.Choice(panel, choices=pitches)
        self.pitch_choice.SetStringSelection(self.temp_settings.get("tts_pitch", "+0Hz"))
        pitch_sizer.Add(self.pitch_choice, 0)
        sizer.Add(pitch_sizer, 0, wx.ALL, 5)

        panel.SetSizer(sizer)
        return panel

    def create_advanced_panel(self, parent):
        panel = wx.Panel(parent)
        sizer = wx.BoxSizer(wx.VERTICAL)

        dir_sizer = wx.BoxSizer(wx.HORIZONTAL)
        dir_sizer.Add(wx.StaticText(panel, label=_("Default Export Directory:")), 0, wx.CENTER)
        self.export_dir = wx.DirPickerCtrl(panel, path=self.temp_settings.get("export_dir", ""))
        dir_sizer.Add(self.export_dir, 1, wx.EXPAND)
        sizer.Add(dir_sizer, 0, wx.ALL | wx.EXPAND, 5)

        vol_sizer = wx.BoxSizer(wx.HORIZONTAL)
        vol_sizer.Add(wx.StaticText(panel, label=_("TTS Volume (0-100):")), 0, wx.CENTER)
        self.vol_spin = wx.SpinCtrl(panel, value=str(self.temp_settings.get("tts_volume", 100)), min=0, max=100)
        vol_sizer.Add(self.vol_spin, 0)
        sizer.Add(vol_sizer, 0, wx.ALL, 5)

        self.chunk_cb = wx.CheckBox(panel, label=_("Enable TTS Chunking (for large texts)"))
        self.chunk_cb.SetValue(self.temp_settings.get("tts_chunking", True))
        sizer.Add(self.chunk_cb, 0, wx.ALL, 5)

        panel.SetSizer(sizer)
        return panel

    def update_tts_voice_list(self):
        voices = self.parent.saved_voices if hasattr(self.parent, 'saved_voices') else []
        if not voices:
            voices = ["en-US-AriaNeural", "en-US-JennyNeural", "en-US-GuyNeural", "en-GB-SoniaNeural"]
        self.voice_choice.SetItems(voices)
        if self.temp_settings["voice"] not in voices:
            self.voice_choice.SetSelection(0)
        else:
            self.voice_choice.SetStringSelection(self.temp_settings["voice"])

    def on_fetch_voices(self, event):
        if edge_tts is None:
            wx.MessageBox(_("edge_tts library not installed or missing dependencies."), _("Error"), wx.OK | wx.ICON_ERROR)
            return
        pd = ProcessingDialog(self, _("Fetching voices..."))
        pd.Show()
        def fetch():
            try:
                async def get_voices():
                    return await edge_tts.list_voices()
                loop = asyncio.new_event_loop()
                asyncio.set_event_loop(loop)
                voices = loop.run_until_complete(get_voices())
                loop.close()
                short_names = [v["ShortName"] for v in voices]
                wx.CallAfter(self.voice_choice.SetItems, short_names)
                if self.temp_settings["voice"] not in short_names:
                    wx.CallAfter(self.voice_choice.SetSelection, 0)
                else:
                    wx.CallAfter(self.voice_choice.SetStringSelection, self.temp_settings["voice"])
                self.parent.saved_voices = short_names
                data = load_data()
                data["saved_voices"] = short_names
                save_data(data)
                thread_safe_msg(_("Voices updated."), _("Success"))
            except Exception as e:
                thread_safe_msg(_("Failed to fetch voices: {}").format(str(e)), _("Error"), True)
            finally:
                wx.CallAfter(pd.Destroy)
        threading.Thread(target=fetch, daemon=True).start()

    def get_settings(self):
        return {
            "voice": self.voice_choice.GetStringSelection(),
            "tts_speed": self.speed_choice.GetStringSelection(),
            "tts_pitch": self.pitch_choice.GetStringSelection(),
            "zoom": self.zoom_spin.GetValue(),
            "contrast": self.contrast_cb.IsChecked(),
            "invert": self.invert_cb.IsChecked(),
            "font_size": self.font_spin.GetValue(),
            "reading_speed": self.speed_spin.GetValue(),
            "extraction_profile": self.profile_choice.GetStringSelection(),
            "export_dir": self.export_dir.GetPath(),
            "tts_volume": self.vol_spin.GetValue(),
            "tts_chunking": self.chunk_cb.IsChecked(),
        }

    def on_ok(self, event):
        self.EndModal(wx.ID_OK)

    def on_cancel(self, event):
        self.EndModal(wx.ID_CANCEL)

class MainDialog(wx.Dialog):
    def __init__(self, parent):
        super().__init__(parent, title=_("Document Reader"))
        self.SetSize((400, 350))
        self.Centre()
        main_sizer = wx.BoxSizer(wx.VERTICAL)

        title_label = wx.StaticText(self, label=_("Document Reader for NVDA"))
        title_font = wx.Font(14, wx.FONTFAMILY_DEFAULT, wx.FONTSTYLE_NORMAL, wx.FONTWEIGHT_BOLD)
        title_label.SetFont(title_font)
        main_sizer.Add(title_label, 0, wx.ALL | wx.CENTER, 15)

        import_btn = wx.Button(self, label=_("&Import Document"))
        import_btn.Bind(wx.EVT_BUTTON, self.on_import)
        main_sizer.Add(import_btn, 0, wx.ALL | wx.EXPAND, 10)

        resume_btn = wx.Button(self, label=_("&Resume Last Document"))
        resume_btn.Bind(wx.EVT_BUTTON, self.on_resume_last)
        main_sizer.Add(resume_btn, 0, wx.ALL | wx.EXPAND, 10)

        recent_btn = wx.Button(self, label=_("Recent &Files"))
        recent_btn.Bind(wx.EVT_BUTTON, self.on_recent_files)
        main_sizer.Add(recent_btn, 0, wx.ALL | wx.EXPAND, 10)

        search_btn = wx.Button(self, label=_("&Search in Last Document"))
        search_btn.Bind(wx.EVT_BUTTON, self.on_search_last)
        main_sizer.Add(search_btn, 0, wx.ALL | wx.EXPAND, 10)

        about_btn = wx.Button(self, label=_("&About"))
        about_btn.Bind(wx.EVT_BUTTON, self.on_about)
        main_sizer.Add(about_btn, 0, wx.ALL | wx.EXPAND, 10)

        help_btn = wx.Button(self, label=_("&Help"))
        help_btn.Bind(wx.EVT_BUTTON, self.on_help)
        main_sizer.Add(help_btn, 0, wx.ALL | wx.EXPAND, 10)

        close_btn = wx.Button(self, label=_("&Close"))
        close_btn.Bind(wx.EVT_BUTTON, self.on_close)
        main_sizer.Add(close_btn, 0, wx.ALL | wx.EXPAND, 10)

        self.SetSizer(main_sizer)
        import_btn.SetFocus()

    def on_import(self, event):
        wildcard = "Supported files (*.pdf;*.docx;*.epub)|*.pdf;*.docx;*.epub|PDF files (*.pdf)|*.pdf|Word files (*.docx)|*.docx|EPUB files (*.epub)|*.epub"
        with wx.FileDialog(self, _("Choose Document"), wildcard=wildcard) as file_dialog:
            if file_dialog.ShowModal() == wx.ID_OK:
                file_path = file_dialog.GetPath()
                self.open_document(file_path)

    def open_document(self, file_path):
        viewer = DocumentViewerDialog(self, file_path)
        viewer.Show()

    def on_resume_last(self, event):
        data = load_data()
        history = data.get("history", {})
        if history:
            last_file = max(history.items(), key=lambda x: x[1])
            file_path = last_file[0]
            if os.path.exists(file_path):
                self.open_document(file_path)
            else:
                show_msg(_("Last document file no longer exists."), _("Error"), True)
        else:
            show_msg(_("No recent document found."), _("Information"))

    def on_recent_files(self, event):
        data = load_data()
        recent = data.get("recent_files", {})
        files = [k for k, v in sorted(recent.items(), key=lambda x: x[1], reverse=True) if os.path.exists(k)][:15]
        if not files:
            show_msg(_("No recent files found."), _("Information"))
            return
        dlg = wx.SingleChoiceDialog(self, _("Select a recent document:"), _("Recent Files"), files)
        if dlg.ShowModal() == wx.ID_OK:
            selected = dlg.GetStringSelection()
            dlg.Destroy()
            if selected:
                self.open_document(selected)
        else:
            dlg.Destroy()

    def on_search_last(self, event):
        data = load_data()
        history = data.get("history", {})
        if history:
            last_file = max(history.items(), key=lambda x: x[1])[0]
            if os.path.exists(last_file):
                self.open_document(last_file)
                return
        show_msg(_("No recent document to search."), _("Information"))

    def on_about(self, event):
        dlg = AboutDialog(self)
        dlg.ShowModal()
        dlg.Destroy()

    def on_help(self, event):
        dlg = HelpDialog(self)
        dlg.ShowModal()
        dlg.Destroy()

    def on_close(self, event):
        self.Close()

class GlobalPlugin(globalPluginHandler.GlobalPlugin):
    def __init__(self):
        super().__init__()
        if globalVars.appArgs.secure:
            return
        self.main_dialog = None
        self.create_menu()

    def create_menu(self):
        self.tools_menu = gui.mainFrame.sysTrayIcon.toolsMenu
        self.doc_reader_item = self.tools_menu.Append(
            wx.ID_ANY,
            _("Document &Reader"),
            _("Open Document Reader")
        )
        gui.mainFrame.sysTrayIcon.Bind(
            wx.EVT_MENU,
            self.on_tools_menu_doc_reader,
            self.doc_reader_item
        )

    def on_tools_menu_doc_reader(self, event):
        self.script_show_main_dialog(None)

    @script(
        description=_("Open Document Reader"),
        category=_("Document Reader"),
        gesture="kb:NVDA+alt+p"
    )
    def script_show_main_dialog(self, gesture):
        if self.main_dialog:
            self.main_dialog.Raise()
            return
        gui.mainFrame.prePopup()
        self.main_dialog = MainDialog(gui.mainFrame)
        self.main_dialog.Show()
        self.main_dialog.Bind(wx.EVT_CLOSE, self.on_main_dialog_close)
        gui.mainFrame.postPopup()

    def on_main_dialog_close(self, event):
        if self.main_dialog:
            self.main_dialog.Destroy()
            self.main_dialog = None
        gui.mainFrame.postPopup()

    def terminate(self):
        if self.main_dialog:
            self.main_dialog.Destroy()
            self.main_dialog = None
        try:
            if self.doc_reader_item:
                self.tools_menu.Remove(self.doc_reader_item)
        except:
            pass
